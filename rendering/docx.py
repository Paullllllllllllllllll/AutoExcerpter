"""DOCX summary document creation for AutoExcerpter.

Generates formatted Word documents with structured summaries, including:
- LaTeX formula parsing and native Word equation rendering (OMML)
- Citation deduplication with OpenAlex hyperlinks
- Document structure overview with page ranges
- XML sanitization for safe DOCX output
"""

from __future__ import annotations

import re
from collections.abc import Callable
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import mathml2omml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from latex2mathml.converter import convert as latex_to_mathml

from config import app as config
from config.constants import (
    BODY_FONT_NAME,
    BODY_FONT_PT,
    BODY_SPACE_AFTER_PT,
    BULLET_FONT_PT,
    BULLET_HANGING_INDENT_CM,
    BULLET_LEFT_INDENT_CM,
    BULLET_SPACE_AFTER_PT,
    COLOR_BLACK,
    COLOR_METADATA_GRAY,
    COLOR_PAGE_HEADING,
    COLOR_REF_META_GRAY,
    COLOR_SECTION_RULE,
    FOOTER_FONT_PT,
    MATH_NAMESPACE,
    METADATA_FONT_PT,
    PAGE_HEADING_FONT_PT,
    PAGE_HEADING_SPACE_AFTER_PT,
    PAGE_HEADING_SPACE_BEFORE_PT,
    PAGE_HEIGHT_CM,
    PAGE_ITEM_HEADING_LEVEL,
    PAGE_MARGIN_CM,
    PAGE_WIDTH_CM,
    REF_FONT_PT,
    REF_HANGING_INDENT_CM,
    REF_META_FONT_PT,
    REF_SPACE_AFTER_PT,
    SECTION_HEADING_FONT_PT,
    SECTION_HEADING_LEVEL,
    SECTION_HEADING_SPACE_AFTER_PT,
    SECTION_HEADING_SPACE_BEFORE_PT,
    TITLE_FONT_PT,
    TITLE_HEADING_LEVEL,
    TITLE_SPACE_AFTER_PT,
)
from config.logger import setup_logger
from rendering.citations import CitationManager, enrich_if_enabled
from rendering.summary import (
    PAGE_TYPE_LABELS,
    STRUCTURE_PAGE_TYPE_ORDER,
    format_structure_page_range,
    prepare_summary_data,
    sanitize_for_xml,
)

logger = setup_logger(__name__)


def parse_latex_in_text(text: str) -> list[tuple[str, str]]:
    """Parse text and split it into segments of regular text and LaTeX formulas.

    Handles:
    - Display math: $$...$$ (processed first to avoid conflicts)
    - Inline math: $...$ (single dollar, non-greedy)
    - Escaped dollar signs: \\$ are preserved as literal $
    - Multi-line formulas (DOTALL mode)
    - Nested braces within formulas

    Returns:
        List of (content, type) tuples where type is 'text', 'latex_display', or
        'latex_inline'.
    """
    if not text:
        return [("", "text")]

    # Placeholder for escaped dollar signs to protect them during parsing
    ESCAPED_DOLLAR_PLACEHOLDER = "\x00ESCAPED_DOLLAR\x00"

    # Protect escaped dollar signs
    protected_text = text.replace("\\$", ESCAPED_DOLLAR_PLACEHOLDER)

    segments: list[tuple[str, str]] = []

    # Pattern for display math $$...$$ (process first — greedy on delimiters,
    # non-greedy on content)
    display_pattern = r"\$\$(.+?)\$\$"

    # Pattern for inline math $...$ (single $, non-greedy, must not be empty)
    inline_pattern = r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)"

    # First pass: extract display math $$...$$
    temp_segments: list[tuple[str, str, int, int]] = []  # (content, type, start, end)

    for match in re.finditer(display_pattern, protected_text, re.DOTALL):
        temp_segments.append(
            (match.group(1), "latex_display", match.start(), match.end())
        )

    # Second pass: extract inline math $...$ from non-display regions
    display_ranges = [(s[2], s[3]) for s in temp_segments]

    def overlaps_display_range(start: int, end: int) -> bool:
        # True if [start, end) intersects any display range, including the case
        # where the inline match fully encloses a $$...$$ block (both endpoints
        # outside the range). Half-open interval overlap test.
        return any(start < d_end and d_start < end for d_start, d_end in display_ranges)

    for match in re.finditer(inline_pattern, protected_text, re.DOTALL):
        if not overlaps_display_range(match.start(), match.end()):
            temp_segments.append(
                (match.group(1), "latex_inline", match.start(), match.end())
            )

    # Sort all segments by position
    temp_segments.sort(key=lambda x: x[2])

    # Build final segments list with text between formulas
    last_end = 0
    for content, seg_type, start, end in temp_segments:
        if start > last_end:
            text_segment = protected_text[last_end:start]
            text_segment = text_segment.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
            if text_segment:
                segments.append((text_segment, "text"))

        restored_content = content.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
        segments.append((restored_content, seg_type))
        last_end = end

    # Add remaining text after last formula
    if last_end < len(protected_text):
        remaining = protected_text[last_end:]
        remaining = remaining.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
        if remaining:
            segments.append((remaining, "text"))

    if not segments:
        return [(text, "text")]

    return segments


def normalize_latex_whitespace(latex_code: str) -> str:
    """Normalize whitespace in LaTeX formulas for consistent processing."""
    normalized = latex_code.strip()
    normalized = re.sub(r" +", " ", normalized)
    normalized = re.sub(r"\s*([=+\-*/^_{}])\s*", r"\1", normalized)
    return normalized


def _regex_rule(
    pattern: str, replacement: str, desc: str, flags: int = 0
) -> tuple[
    Callable[[str], bool],
    Callable[[str], str],
    str,
]:
    """Create a (should_apply, transform, description) tuple for regex-based rules."""
    compiled = re.compile(pattern, flags)

    def should_apply(text: str) -> bool:
        return bool(compiled.search(text))

    def transform(text: str) -> str:
        return compiled.sub(replacement, text)

    return should_apply, transform, desc


def _literal_rule(
    old: str, new: str, desc: str
) -> tuple[
    Callable[[str], bool],
    Callable[[str], str],
    str,
]:
    """Create a (should_apply, transform, description) tuple for literal string
    rules."""

    def should_apply(text: str) -> bool:
        return old in text

    def transform(text: str) -> str:
        return text.replace(old, new)

    return should_apply, transform, desc


def _simplify_delimiter_sizing(text: str) -> str:
    """Replace all delimiter size commands with plain delimiters."""
    sizes = [
        "\\Big",
        "\\big",
        "\\bigg",
        "\\Bigg",
        "\\Bigl",
        "\\Bigr",
        "\\bigl",
        "\\bigr",
        "\\biggl",
        "\\biggr",
        "\\Biggl",
        "\\Biggr",
        "\\bigm",
        "\\Bigm",
    ]
    delims = ["(", ")", "[", "]", "|", "\\{", "\\}", "\\|", "."]
    for size in sizes:
        for delim in delims:
            target = f"{size}{delim}"
            if target in text:
                plain = delim if delim not in ("\\{", "\\}", "\\|") else delim[1:]
                text = text.replace(target, plain if delim != "." else "")
    return text


def _simplify_auto_sizing(text: str) -> str:
    """Replace \\left/\\right auto-sizing delimiters."""
    text = text.replace("\\left.", "")
    text = text.replace("\\right.", "")
    text = re.sub(r"\\left\s*([(\[|])", r"\1", text)
    text = re.sub(r"\\right\s*([)\]|])", r"\1", text)
    text = re.sub(r"\\left\s*\\([{|}])", r"\\\1", text)
    text = re.sub(r"\\right\s*\\([{|}])", r"\\\1", text)
    text = text.replace("\\middle|", "|")
    text = text.replace("\\middle", "")
    return text


type _LaTeXRule = tuple[Callable[[str], bool], Callable[[str], str], str]

# Module-level list: each entry is (should_apply, transform, description)
_LATEX_SIMPLIFICATIONS: list[_LaTeXRule] = [
    # Logical operators
    _literal_rule("\\land", "\\wedge", "logical operator (\\land -> \\wedge)"),
    _literal_rule("\\lor", "\\vee", "logical operator (\\lor -> \\vee)"),
    _literal_rule("\\lnot", "\\neg", "logical operator (\\lnot -> \\neg)"),
    _literal_rule(
        "\\iff", "\\Leftrightarrow", "logical operator (\\iff -> \\Leftrightarrow)"
    ),
    _literal_rule(
        "\\implies", "\\Rightarrow", "logical operator (\\implies -> \\Rightarrow)"
    ),
    # Text commands
    _regex_rule(r"\\text\{([^}]*)\}", r"\\mathrm{\1}", "text -> mathrm"),
    _regex_rule(r"\\textrm\{([^}]*)\}", r"\\mathrm{\1}", "textrm -> mathrm"),
    _regex_rule(r"\\textit\{([^}]*)\}", r"\\mathit{\1}", "textit -> mathit"),
    _regex_rule(r"\\textbf\{([^}]*)\}", r"\\mathbf{\1}", "textbf -> mathbf"),
    # Accent macros
    _regex_rule(r"\\bar\{([^}]+)\}", r"\\overline{\1}", "accent (bar -> overline)"),
    _regex_rule(
        r"\\tilde\{([^}]+)\}", r"\\widetilde{\1}", "accent (tilde -> widetilde)"
    ),
    _regex_rule(r"\\hat\{([^}]+)\}", r"\\widehat{\1}", "accent (hat -> widehat)"),
    _regex_rule(
        r"\\vec\{([^}]+)\}", r"\\overrightarrow{\1}", "accent (vec -> overrightarrow)"
    ),
    _regex_rule(r"\\dot\{([^}]+)\}", r"\1", "accent (dot removed)"),
    _regex_rule(r"\\ddot\{([^}]+)\}", r"\1", "accent (ddot removed)"),
    # Delimiter sizing (complex transform)
    (
        lambda text: any(
            f"{s}{d}" in text
            for s in [
                "\\Big",
                "\\big",
                "\\bigg",
                "\\Bigg",
                "\\Bigl",
                "\\Bigr",
                "\\bigl",
                "\\bigr",
                "\\biggl",
                "\\biggr",
                "\\Biggl",
                "\\Biggr",
                "\\bigm",
                "\\Bigm",
            ]
            for d in ["(", ")", "[", "]", "|", "\\{", "\\}", "\\|", "."]
        ),
        _simplify_delimiter_sizing,
        "delimiter sizing removed",
    ),
    # Auto-sizing delimiters
    (
        lambda text: "\\left" in text or "\\right" in text,
        _simplify_auto_sizing,
        "auto-sizing delimiters (left/right) removed",
    ),
    # Overbrace/underbrace
    _regex_rule(r"\\overbrace\{([^}]+)\}\^?\{?[^}]*\}?", r"\1", "overbrace simplified"),
    _regex_rule(
        r"\\underbrace\{([^}]+)\}_?\{?[^}]*\}?", r"\1", "underbrace simplified"
    ),
    _regex_rule(r"\\overleftarrow\{([^}]+)\}", r"\1", "overleftarrow simplified"),
    _regex_rule(
        r"\\overrightarrow\{([^}]+)\}", r"\\vec{\1}", "overrightarrow simplified"
    ),
    _regex_rule(r"\\underleftarrow\{([^}]+)\}", r"\1", "underleftarrow simplified"),
    _regex_rule(r"\\underrightarrow\{([^}]+)\}", r"\1", "underrightarrow simplified"),
    # Phantoms
    _regex_rule(r"\\phantom\{[^}]*\}", "", "phantom removed"),
    _regex_rule(r"\\hphantom\{[^}]*\}", "", "hphantom removed"),
    _regex_rule(r"\\vphantom\{[^}]*\}", "", "vphantom removed"),
    # Spacing literals
    _literal_rule("\\,", " ", "thin space normalized"),
    _literal_rule("\\;", " ", "thick space normalized"),
    _literal_rule("\\!", "", "negative thin space normalized"),
    _literal_rule("\\quad", " ", "quad normalized"),
    _literal_rule("\\qquad", "  ", "qquad normalized"),
    # Spacing patterns
    _regex_rule(r"\\hspace\{[^}]*\}", " ", "hspace normalized"),
    _regex_rule(r"\\vspace\{[^}]*\}", "", "vspace normalized"),
    _regex_rule(r"\\mspace\{[^}]*\}", " ", "mspace normalized"),
    # Operatorname
    _regex_rule(
        r"\\operatorname\{([^}]+)\}", r"\\mathrm{\1}", "operatorname -> mathrm"
    ),
    # Environments
    _regex_rule(
        r"\\begin\{align\*?\}(.+?)\\end\{align\*?\}",
        r"\1",
        "align env unwrapped",
        re.DOTALL,
    ),
    _regex_rule(
        r"\\begin\{gather\*?\}(.+?)\\end\{gather\*?\}",
        r"\1",
        "gather env unwrapped",
        re.DOTALL,
    ),
    _regex_rule(
        r"\\begin\{equation\*?\}(.+?)\\end\{equation\*?\}",
        r"\1",
        "equation env unwrapped",
        re.DOTALL,
    ),
    _regex_rule(
        r"\\begin\{split\}(.+?)\\end\{split\}", r"\1", "split env unwrapped", re.DOTALL
    ),
    _regex_rule(
        r"\\begin\{cases\}(.+?)\\end\{cases\}", r"\1", "cases env unwrapped", re.DOTALL
    ),
    # Alignment markers
    (
        lambda text: "&" in text or "\\\\" in text,
        lambda text: text.replace("&", " ").replace("\\\\", " "),
        "alignment markers cleaned",
    ),
    # Limits modifiers
    (
        lambda text: "\\limits" in text or "\\nolimits" in text,
        lambda text: text.replace("\\limits", "").replace("\\nolimits", ""),
        "limits modifiers removed",
    ),
]


def simplify_problematic_latex(latex_code: str) -> tuple[str, list[str]]:
    """Simplify LaTeX code by removing constructs known to cause mathml2omml issues.

    Returns:
        Tuple of (simplified_latex, list_of_applied_simplifications)
    """
    simplified = latex_code
    applied: list[str] = []

    # Normalize whitespace first
    original_len = len(simplified)
    simplified = normalize_latex_whitespace(simplified)
    if len(simplified) != original_len:
        applied.append("whitespace normalization")

    # Apply data-driven rules
    for should_apply, transform, desc in _LATEX_SIMPLIFICATIONS:
        if should_apply(simplified):
            simplified = transform(simplified)
            applied.append(desc)

    # Final whitespace cleanup
    simplified = re.sub(r"\s+", " ", simplified).strip()

    return simplified, applied


def sanitize_omml_xml(omml_markup: str) -> str:
    """Attempt to fix common XML issues in OMML markup generated by mathml2omml."""
    try:
        ET.fromstring(omml_markup)
        return omml_markup
    except ET.ParseError as e:
        error_msg = str(e)

        if "groupChr" in error_msg:
            fixed_markup = omml_markup

            pattern = r"<m:groupChrPr[^>]*>"
            positions = [
                (m.start(), m.end()) for m in re.finditer(pattern, fixed_markup)
            ]

            for _start, end in reversed(positions):
                after_tag = fixed_markup[end:]
                close_prop = after_tag.find("</m:groupChrPr>")
                close_group = after_tag.find("</m:groupChr>")

                if close_group != -1 and (close_prop == -1 or close_prop > close_group):
                    insert_pos = end + close_group
                    fixed_markup = (
                        fixed_markup[:insert_pos]
                        + "</m:groupChrPr>"
                        + fixed_markup[insert_pos:]
                    )

            try:
                ET.fromstring(fixed_markup)
                logger.info("Successfully sanitized OMML with groupChr fix")
                return fixed_markup
            except ET.ParseError as parse_err:
                logger.debug("First sanitization attempt failed: %s", parse_err)

                fixed_markup2 = omml_markup
                max_iterations = 10
                for _ in range(max_iterations):
                    match = re.search(
                        r"<m:groupChr>\s*<m:groupChrPr[^>]*>(?:(?!</m:groupChrPr>).)*?</m:groupChr>",
                        fixed_markup2,
                        flags=re.DOTALL,
                    )
                    if not match:
                        break
                    block = match.group(0)
                    fixed_block = block.replace(
                        "</m:groupChr>", "</m:groupChrPr></m:groupChr>", 1
                    )
                    fixed_markup2 = (
                        fixed_markup2[: match.start()]
                        + fixed_block
                        + fixed_markup2[match.end() :]
                    )

                try:
                    ET.fromstring(fixed_markup2)
                    logger.info(
                        "Successfully sanitized OMML with iterative groupChr fix"
                    )
                    return fixed_markup2
                except ET.ParseError:
                    logger.debug("Iterative sanitization failed, trying fallback")

                    fixed_markup3 = re.sub(
                        r"<m:groupChr>\s*<m:groupChrPr[^>]*>.*?</m:groupChr>",
                        lambda m: (
                            m.group(0)
                            .replace("<m:groupChrPr", "<!--m:groupChrPr")
                            .replace("</m:groupChrPr>", "-->")
                        ),
                        omml_markup,
                        flags=re.DOTALL,
                    )

                    try:
                        ET.fromstring(fixed_markup3)
                        logger.info(
                            "Successfully sanitized OMML by commenting out groupChrPr"
                        )
                        return fixed_markup3
                    except ET.ParseError:
                        pass

        return omml_markup


def _ensure_omml_namespace(omml_markup: str) -> str:
    """Ensure the OMML markup carries the math namespace and an <m:oMath> root."""
    if "xmlns:m" not in omml_markup.split("\n", 1)[0]:
        omml_markup = omml_markup.replace(
            "<m:oMath",
            f'<m:oMath xmlns:m="{MATH_NAMESPACE}"',
            1,
        )

    if not omml_markup.strip().startswith("<m:oMath"):
        omml_markup = f'<m:oMath xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMath>'

    return omml_markup


def add_math_to_paragraph(paragraph: Any, latex_code: str) -> None:
    """Add a mathematical formula to a paragraph using OMML for native Word equation
    rendering."""
    try:
        mathml = latex_to_mathml(latex_code)
        omml_markup = _ensure_omml_namespace(mathml2omml.convert(mathml))

        omml_para = (
            f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMathPara>'
        )

        try:
            omml_element = parse_xml(omml_para)
        except (ET.ParseError, ValueError) as parse_error:
            logger.warning(
                "XML parsing failed, attempting to sanitize OMML: %s", parse_error
            )
            sanitized_omml = sanitize_omml_xml(omml_markup)
            sanitized_para = (
                f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">'
                f"{sanitized_omml}</m:oMathPara>"
            )
            omml_element = parse_xml(sanitized_para)

        paragraph._p.append(omml_element)

    except Exception as exc:
        try:
            simplified_latex, simplifications = simplify_problematic_latex(latex_code)

            if simplified_latex != latex_code:
                logger.info(
                    "Attempting conversion with simplified LaTeX: %s",
                    ", ".join(simplifications),
                )
                mathml = latex_to_mathml(simplified_latex)
                omml_markup = _ensure_omml_namespace(mathml2omml.convert(mathml))

                omml_para = (
                    f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">'
                    f"{omml_markup}</m:oMathPara>"
                )
                omml_element = parse_xml(omml_para)
                paragraph._p.append(omml_element)
                logger.info("Successfully converted simplified LaTeX")
                return
        except Exception:
            pass

        logger.warning(
            "Failed to convert LaTeX to MathML: %s. Displaying as text.", exc
        )
        logger.warning("Problematic LaTeX code: %s", latex_code[:200])
        run = paragraph.add_run(f" {sanitize_for_xml(latex_code)} ")
        run.font.name = "Cambria Math"
        run.italic = True


def add_formatted_text_to_paragraph(paragraph: Any, text: str) -> None:
    """Add text to a paragraph, parsing and rendering LaTeX formulas."""
    segments = parse_latex_in_text(text)

    for content, segment_type in segments:
        if segment_type in ("latex", "latex_display", "latex_inline"):
            add_math_to_paragraph(paragraph, content)
        else:
            paragraph.add_run(sanitize_for_xml(content))


_MARKDOWN_EMPHASIS_PATTERN = re.compile(
    r"\*\*(?P<bold>[^*]+)\*\*|\*(?P<italic>[^*]+)\*"
)


def parse_markdown_emphasis(text: str) -> list[tuple[str, str]]:
    """Split text into ('text' | 'bold' | 'italic', content) segments.

    Citation strings extracted by the LLM occasionally carry Markdown emphasis
    (e.g. ``*The American Economic Review, 87*(2)``). The Markdown writer
    renders these natively; the DOCX writer maps them to run formatting.
    """
    segments: list[tuple[str, str]] = []
    last_end = 0
    for match in _MARKDOWN_EMPHASIS_PATTERN.finditer(text):
        if match.start() > last_end:
            segments.append(("text", text[last_end : match.start()]))
        if match.group("bold") is not None:
            segments.append(("bold", match.group("bold")))
        else:
            segments.append(("italic", match.group("italic")))
        last_end = match.end()
    if last_end < len(text):
        segments.append(("text", text[last_end:]))
    return segments or [("text", text)]


def strip_markdown_emphasis(text: str) -> str:
    """Remove Markdown emphasis markers, keeping the emphasized text."""
    return _MARKDOWN_EMPHASIS_PATTERN.sub(
        lambda m: m.group("bold") or m.group("italic"), text
    )


def add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Add a hyperlink to a paragraph in a DOCX document."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _rgb(color: int) -> RGBColor:
    """Build an :class:`RGBColor` from a ``0xRRGGBB`` integer."""
    return RGBColor((color >> 16) & 0xFF, (color >> 8) & 0xFF, color & 0xFF)


def _set_east_asian_font(element: Any, name: str) -> None:
    """Pin the East Asian font so Word does not substitute a different face.

    Also strips the theme-font attributes that the built-in Title/Heading styles
    carry on their ``w:rFonts`` element. Word and LibreOffice give ``*Theme``
    attributes precedence over an explicit ``w:ascii``/``w:hAnsi`` face, so
    without this the headings would render in the theme's major font (Calibri/
    Carlito) instead of the requested body face. ``w:cs`` is set for
    completeness so complex-script runs match too.
    """
    rpr = element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)
    for theme_attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{theme_attr}")
        if rfonts.get(key) is not None:
            del rfonts.attrib[key]


def _style_font(
    style: Any,
    *,
    size_pt: float,
    bold: bool | None = None,
    color: int | None = None,
    name: str = BODY_FONT_NAME,
) -> None:
    """Apply an explicit font to a named style (never rely on docx defaults)."""
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = _rgb(color)
        # Strip inherited theme colors so the explicit RGB always wins in Word.
        rpr = style.element.get_or_add_rPr()
        color_el = rpr.find(qn("w:color"))
        if color_el is not None:
            for attr in ("themeColor", "themeTint", "themeShade"):
                key = qn(f"w:{attr}")
                if color_el.get(key) is not None:
                    del color_el.attrib[key]
    _set_east_asian_font(style.element, name)


def _clear_style_borders(style: Any) -> None:
    """Remove any paragraph border defined on a style (e.g. Title underline)."""
    ppr = style.element.get_or_add_pPr()
    for existing in ppr.findall(qn("w:pBdr")):
        ppr.remove(existing)


def _set_style_bottom_border(style: Any, color: int) -> None:
    """Give a style a thin gray bottom border used as a visual section rule."""
    ppr = style.element.get_or_add_pPr()
    for existing in ppr.findall(qn("w:pBdr")):
        ppr.remove(existing)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")  # eighths of a point -> ~0.5 pt
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), f"{color:06X}")
    pbdr.append(bottom)
    ppr.append(pbdr)


def _apply_document_styles(document: Any) -> None:
    """Define the summary document's typography once, on named styles.

    Sets Normal (body), Title, Heading 1 (section headings), Heading 2 (per-page
    headings), and List Bullet so downstream paragraphs need no per-run
    formatting. Concrete sizes, colors, and spacing live in ``config.constants``.
    """
    styles = document.styles

    normal = styles["Normal"]
    _style_font(normal, size_pt=BODY_FONT_PT, bold=False, color=COLOR_BLACK)
    npf = normal.paragraph_format
    npf.space_before = Pt(0)
    npf.space_after = Pt(BODY_SPACE_AFTER_PT)
    npf.line_spacing = 1.0
    npf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    title = styles["Title"]
    _style_font(title, size_pt=TITLE_FONT_PT, bold=True, color=COLOR_BLACK)
    _clear_style_borders(title)
    tpf = title.paragraph_format
    tpf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    tpf.space_before = Pt(0)
    tpf.space_after = Pt(TITLE_SPACE_AFTER_PT)

    section_heading = styles["Heading 1"]
    _style_font(
        section_heading,
        size_pt=SECTION_HEADING_FONT_PT,
        bold=True,
        color=COLOR_BLACK,
    )
    _set_style_bottom_border(section_heading, COLOR_SECTION_RULE)
    shp = section_heading.paragraph_format
    shp.space_before = Pt(SECTION_HEADING_SPACE_BEFORE_PT)
    shp.space_after = Pt(SECTION_HEADING_SPACE_AFTER_PT)
    shp.keep_with_next = True

    page_heading = styles["Heading 2"]
    _style_font(
        page_heading,
        size_pt=PAGE_HEADING_FONT_PT,
        bold=True,
        color=COLOR_PAGE_HEADING,
    )
    php = page_heading.paragraph_format
    php.space_before = Pt(PAGE_HEADING_SPACE_BEFORE_PT)
    php.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)
    php.keep_with_next = True

    bullet = styles["List Bullet"]
    _style_font(bullet, size_pt=BULLET_FONT_PT, color=COLOR_BLACK)
    bpf = bullet.paragraph_format
    bpf.left_indent = Cm(BULLET_LEFT_INDENT_CM)
    bpf.first_line_indent = -Cm(BULLET_HANGING_INDENT_CM)
    bpf.space_before = Pt(0)
    bpf.space_after = Pt(BULLET_SPACE_AFTER_PT)


def _add_page_number_footer(section: Any) -> None:
    """Add a bottom-right page-number field to a section footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run()
    run.font.name = BODY_FONT_NAME
    run.font.size = Pt(FOOTER_FONT_PT)
    _set_east_asian_font(run._r, BODY_FONT_NAME)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _configure_page(document: Any) -> None:
    """Apply A4 geometry, uniform margins, and a page-number footer."""
    for section in document.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(PAGE_MARGIN_CM)
        section.bottom_margin = Cm(PAGE_MARGIN_CM)
        section.left_margin = Cm(PAGE_MARGIN_CM)
        section.right_margin = Cm(PAGE_MARGIN_CM)
        _add_page_number_footer(section)


def create_docx_summary(
    summary_results: list[dict[str, Any]],
    output_path: Path,
    document_name: str,
    citation_manager: CitationManager | None = None,
    data: Any = None,
) -> None:
    """Create a DOCX summary document from structured summary results.

    Output order:
    1. Title + Metadata
    2. Document Structure (page type overview)
    3. Content Summaries (in document order)
    4. Consolidated References

    When *citation_manager* and *data* are supplied by the pipeline they are
    already consolidated and OpenAlex-enriched, so this writer renders them as-is
    (one enrichment pass per item, shared with the Markdown writer). When called
    standalone they are built and enriched here.
    """
    if citation_manager is None or data is None:
        citation_manager = CitationManager(
            polite_pool_email=config.CITATION_OPENALEX_EMAIL
        )
        data = prepare_summary_data(summary_results, citation_manager)
        citation_manager.consolidate()
        enrich_if_enabled(citation_manager)
    filtered_results = data.filtered_results
    page_type_pages = data.page_type_pages

    document = Document()
    _apply_document_styles(document)
    _configure_page(document)

    # === SECTION 1: Title + metadata ===
    # The document name is the title (prominent); the metadata line records that
    # this is a generated summary.
    title = document.add_heading(sanitize_for_xml(document_name), TITLE_HEADING_LEVEL)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    metadata = (
        f"Processed {datetime.now():%Y-%m-%d %H:%M}"
        f" | {data.content_page_count} content pages"
        f" | {len(filtered_results)} total pages"
    )
    meta_paragraph = document.add_paragraph()
    meta_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_paragraph.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)
    meta_run = meta_paragraph.add_run(metadata)
    meta_run.font.name = BODY_FONT_NAME
    meta_run.font.size = Pt(METADATA_FONT_PT)
    meta_run.font.color.rgb = _rgb(COLOR_METADATA_GRAY)

    # === SECTION 2: Document Structure ===
    has_structure_info = any(pages for pages in page_type_pages.values())
    if has_structure_info:
        document.add_heading("Document Structure", SECTION_HEADING_LEVEL)

        for pt in STRUCTURE_PAGE_TYPE_ORDER:
            pages = page_type_pages.get(pt, [])
            if pages:
                label = PAGE_TYPE_LABELS.get(pt, pt.replace("_", " ").title())
                page_range = format_structure_page_range(pages)
                struct_para = document.add_paragraph()
                label_run = struct_para.add_run(f"{label}: ")
                label_run.bold = True
                struct_para.add_run(page_range)

    # === SECTION 3: Content Summaries ===
    for page_item in data.page_render_items:
        document.add_heading(page_item.heading_text, PAGE_ITEM_HEADING_LEVEL)

        for point in page_item.bullet_points:
            paragraph = document.add_paragraph(style="List Bullet")
            add_formatted_text_to_paragraph(paragraph, point)

    # === SECTION 4: Consolidated References ===
    if citation_manager.citations:
        logger.info(
            "Processing %d unique citations for consolidated references section",
            len(citation_manager.citations),
        )

        document.add_page_break()  # type: ignore[no-untyped-call]
        document.add_heading("Consolidated References", SECTION_HEADING_LEVEL)

        note_text = (
            "The following references were extracted from the document and "
            "consolidated. Duplicate citations have been merged, showing all "
            "pages where each citation appears. Where available, hyperlinks "
            "provide access to extended metadata via OpenAlex."
        )
        note_paragraph = document.add_paragraph()
        note_run = note_paragraph.add_run(note_text)
        note_run.italic = True
        note_run.font.size = Pt(METADATA_FONT_PT)
        note_run.font.color.rgb = _rgb(COLOR_METADATA_GRAY)

        citations_with_pages = citation_manager.get_citations_with_pages()

        for idx, (citation, page_range_str) in enumerate(citations_with_pages, start=1):
            ref_paragraph = document.add_paragraph()
            rpf = ref_paragraph.paragraph_format
            rpf.left_indent = Cm(REF_HANGING_INDENT_CM)
            rpf.first_line_indent = -Cm(REF_HANGING_INDENT_CM)
            rpf.space_after = Pt(REF_SPACE_AFTER_PT)

            num_run = ref_paragraph.add_run(f"[{idx}] ")
            num_run.bold = True
            num_run.font.size = Pt(REF_FONT_PT)

            citation_text = sanitize_for_xml(citation.raw_text)

            if citation.url:
                # A hyperlink is a single styled run; drop emphasis markers.
                add_hyperlink(
                    ref_paragraph, citation.url, strip_markdown_emphasis(citation_text)
                )
            else:
                for kind, content in parse_markdown_emphasis(citation_text):
                    text_run = ref_paragraph.add_run(content)
                    text_run.font.size = Pt(REF_FONT_PT)
                    if kind == "italic":
                        text_run.italic = True
                    elif kind == "bold":
                        text_run.bold = True

            if page_range_str:
                page_run = ref_paragraph.add_run(f" ({page_range_str})")
                page_run.italic = True
                page_run.font.size = Pt(REF_FONT_PT)

            if citation.metadata:
                meta_info_parts = []
                if citation.doi:
                    meta_info_parts.append(f"DOI: {citation.doi}")
                if citation.metadata.get("publication_year"):
                    meta_info_parts.append(
                        f"Year: {citation.metadata['publication_year']}"
                    )

                if meta_info_parts:
                    meta_run = ref_paragraph.add_run(f" [{', '.join(meta_info_parts)}]")
                    meta_run.font.size = Pt(REF_META_FONT_PT)
                    meta_run.italic = True
                    meta_run.font.color.rgb = _rgb(COLOR_REF_META_GRAY)

    document.save(str(output_path))
    logger.info("Summary document saved to %s", output_path)

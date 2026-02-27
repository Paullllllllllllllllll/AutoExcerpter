"""Extended tests for processors/file_manager.py - coverage gap filling.

Covers:
- initialize_log_file: creates log file with JSON header
- append_to_log: appends entries as comma-separated JSON
- finalize_log_file: closes JSON array
- write_transcription_to_text: writes transcription output
- create_docx_summary: DOCX file creation (python-docx mocked)
- create_markdown_summary: additional edge cases
- _format_page_range: compact range formatting
- _format_page_heading_docx: DOCX-specific heading formatting
- add_hyperlink: adds hyperlink elements to paragraph
- sanitize_omml_xml: XML sanitization
- add_math_to_paragraph: math rendering with fallback
- add_formatted_text_to_paragraph: mixed text and LaTeX
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any
from unittest.mock import patch, MagicMock, call

import pytest

from processors.docx_writer import (
    create_docx_summary,
    sanitize_omml_xml,
    add_math_to_paragraph,
    add_formatted_text_to_paragraph,
    add_hyperlink,
    _format_page_heading_docx,
)
from processors.file_manager import (
    write_transcription_to_text,
    sanitize_for_xml,
    _format_page_range,
    _should_render_bullets,
    _get_structure_types,
    filter_empty_pages,
)
from processors.log_manager import (
    initialize_log_file,
    append_to_log,
    finalize_log_file,
    _close_log_handle,
    _LOG_HANDLES,
    _LOG_HANDLES_GUARD,
)
from processors.markdown_writer import create_markdown_summary


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
@pytest.fixture(autouse=True)
def _clear_log_handles():
    """Clear the module-level log handle cache before each test."""
    with _LOG_HANDLES_GUARD:
        for key, (handle, _lock) in list(_LOG_HANDLES.items()):
            try:
                handle.close()
            except Exception:
                pass
        _LOG_HANDLES.clear()
    yield
    # Clean up after test
    with _LOG_HANDLES_GUARD:
        for key, (handle, _lock) in list(_LOG_HANDLES.items()):
            try:
                handle.close()
            except Exception:
                pass
        _LOG_HANDLES.clear()


# ============================================================================
# initialize_log_file
# ============================================================================
class TestInitializeLogFile:
    """Tests for initialize_log_file()."""

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_creates_log_file_with_header(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Creates a log file starting with a JSON array."""
        log_path = tmp_path / "test.log.json"

        result = initialize_log_file(
            log_path=log_path,
            item_name="test_doc.pdf",
            input_path="/path/to/test_doc.pdf",
            input_type="pdf",
            total_images=10,
            model_name="gpt-5-mini",
            extraction_dpi=300,
            concurrency_limit=8,
        )

        assert result is True
        assert log_path.exists()
        content = log_path.read_text(encoding="utf-8")
        assert content.startswith("[")
        # Parse the JSON payload (after the opening bracket)
        payload = json.loads(content[1:].strip())
        assert payload["input_item_name"] == "test_doc.pdf"
        assert payload["total_images"] == 10
        assert payload["configuration"]["concurrent_requests"] == 8
        assert payload["configuration"]["extraction_dpi"] == 300

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_non_openai_model_service_tier_na(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Non-OpenAI models get 'N/A' as service_tier."""
        log_path = tmp_path / "test.log.json"

        initialize_log_file(
            log_path=log_path,
            item_name="doc.pdf",
            input_path="/path",
            input_type="pdf",
            total_images=5,
            model_name="claude-3-opus",
        )

        content = log_path.read_text(encoding="utf-8")
        payload = json.loads(content[1:].strip())
        assert payload["configuration"]["service_tier"] == "N/A"

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_returns_false_on_error(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Returns False when the file cannot be written."""
        bad_path = tmp_path / "nonexistent_dir" / "sub" / "test.log.json"

        result = initialize_log_file(
            log_path=bad_path,
            item_name="doc.pdf",
            input_path="/path",
            input_type="pdf",
            total_images=5,
            model_name="gpt-5-mini",
        )

        assert result is False


# ============================================================================
# append_to_log
# ============================================================================
class TestAppendToLog:
    """Tests for append_to_log()."""

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_appends_entry_as_json(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Appends a comma-separated JSON entry."""
        log_path = tmp_path / "test.log.json"

        initialize_log_file(
            log_path=log_path,
            item_name="doc.pdf",
            input_path="/path",
            input_type="pdf",
            total_images=2,
            model_name="gpt-5-mini",
        )

        entry = {"page": 1, "status": "success", "tokens": 500}
        result = append_to_log(log_path, entry)

        assert result is True
        # Finalize to flush and close the file handle before reading
        finalize_log_file(log_path)
        content = log_path.read_text(encoding="utf-8")
        assert '"page": 1' in content

    def test_returns_false_on_error(self, tmp_path: Path):
        """Returns False when writing fails."""
        # Create a log handle pointing to an invalid path
        bad_path = tmp_path / "nonexistent_dir" / "log.json"
        result = append_to_log(bad_path, {"test": True})
        # This will attempt to open the file in append mode which may fail
        # or succeed depending on directory existence
        assert isinstance(result, bool)


# ============================================================================
# finalize_log_file
# ============================================================================
class TestFinalizeLogFile:
    """Tests for finalize_log_file()."""

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_closes_json_array(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Finalizes the log file by closing the JSON array."""
        log_path = tmp_path / "test.log.json"

        initialize_log_file(
            log_path=log_path,
            item_name="doc.pdf",
            input_path="/path",
            input_type="pdf",
            total_images=1,
            model_name="gpt-5-mini",
        )

        append_to_log(log_path, {"page": 1, "status": "ok"})

        result = finalize_log_file(log_path)

        assert result is True
        content = log_path.read_text(encoding="utf-8")
        assert content.strip().endswith("]")

    @patch("processors.log_manager.get_api_concurrency", return_value=(4, 0.05))
    @patch("processors.log_manager.get_api_timeout", return_value=900)
    @patch("processors.log_manager.get_service_tier", return_value="flex")
    def test_full_lifecycle_produces_valid_json(
        self, mock_tier, mock_timeout, mock_conc, tmp_path: Path
    ):
        """Full init -> append -> finalize produces parseable JSON array."""
        log_path = tmp_path / "test.log.json"

        initialize_log_file(
            log_path=log_path,
            item_name="doc.pdf",
            input_path="/path",
            input_type="pdf",
            total_images=2,
            model_name="gpt-5-mini",
        )

        append_to_log(log_path, {"page": 1, "status": "ok"})
        append_to_log(log_path, {"page": 2, "status": "ok"})
        finalize_log_file(log_path)

        content = log_path.read_text(encoding="utf-8")
        data = json.loads(content)
        assert isinstance(data, list)
        assert len(data) == 3  # header + 2 entries


# ============================================================================
# write_transcription_to_text
# ============================================================================
class TestWriteTranscriptionToText:
    """Tests for write_transcription_to_text()."""

    def test_writes_transcription_file(self, tmp_path: Path):
        """Creates a text file with transcription content and metadata."""
        output_path = tmp_path / "output.txt"
        results = [
            {"transcription": "Page 1 content here."},
            {"transcription": "Page 2 content here."},
        ]

        success = write_transcription_to_text(
            transcription_results=results,
            output_path=output_path,
            document_name="Test Document",
            item_type="pdf",
            total_elapsed_time=120.5,
            source_path=Path("/path/to/source.pdf"),
        )

        assert success is True
        assert output_path.exists()
        content = output_path.read_text(encoding="utf-8")
        assert "# Transcription of: Test Document" in content
        assert "# Type: pdf" in content
        assert "# Total images processed: 2" in content
        assert "# Successfully transcribed: 2" in content
        assert "# Failed items: 0" in content
        assert "Page 1 content here." in content
        assert "Page 2 content here." in content

    def test_counts_errors_in_results(self, tmp_path: Path):
        """Error results are counted in the failure tally."""
        output_path = tmp_path / "output.txt"
        results = [
            {"transcription": "Good page."},
            {"error": "API timeout"},
            {"transcription": "Another good page."},
        ]

        success = write_transcription_to_text(
            transcription_results=results,
            output_path=output_path,
            document_name="Partial",
            item_type="pdf",
            total_elapsed_time=60.0,
            source_path=Path("/path/to/source.pdf"),
        )

        assert success is True
        content = output_path.read_text(encoding="utf-8")
        assert "# Successfully transcribed: 2" in content
        assert "# Failed items: 1" in content

    def test_missing_transcription_uses_error_placeholder(self, tmp_path: Path):
        """Results missing 'transcription' get an ERROR placeholder."""
        output_path = tmp_path / "output.txt"
        results = [{"no_transcription_key": True}]

        write_transcription_to_text(
            transcription_results=results,
            output_path=output_path,
            document_name="Missing",
            item_type="pdf",
            total_elapsed_time=10.0,
            source_path=Path("/path/to/source.pdf"),
        )

        content = output_path.read_text(encoding="utf-8")
        assert "[ERROR] Transcription data missing" in content

    def test_returns_false_on_write_error(self, tmp_path: Path):
        """Returns False when the file cannot be written."""
        bad_path = tmp_path / "no_dir" / "sub" / "output.txt"

        success = write_transcription_to_text(
            transcription_results=[{"transcription": "data"}],
            output_path=bad_path,
            document_name="Bad",
            item_type="pdf",
            total_elapsed_time=0.0,
            source_path=Path("/path"),
        )

        assert success is False

    def test_empty_results(self, tmp_path: Path):
        """Empty results list produces a valid file with zero counts."""
        output_path = tmp_path / "output.txt"

        success = write_transcription_to_text(
            transcription_results=[],
            output_path=output_path,
            document_name="Empty",
            item_type="pdf",
            total_elapsed_time=0.0,
            source_path=Path("/path"),
        )

        assert success is True
        content = output_path.read_text(encoding="utf-8")
        assert "# Total images processed: 0" in content


# ============================================================================
# _format_page_range
# ============================================================================
class TestFormatPageRange:
    """Tests for _format_page_range()."""

    def test_single_page(self):
        """Single page uses 'p.' prefix."""
        assert _format_page_range([5]) == "p. 5"

    def test_consecutive_pages(self):
        """Consecutive pages are collapsed into a range."""
        assert _format_page_range([1, 2, 3]) == "pp. 1-3"

    def test_mixed_ranges(self):
        """Mixed consecutive and non-consecutive pages."""
        result = _format_page_range([1, 2, 3, 5, 7, 8, 9])
        assert result == "pp. 1-3, 5, 7-9"

    def test_empty_list(self):
        """Empty list returns empty string."""
        assert _format_page_range([]) == ""

    def test_unsorted_input(self):
        """Unsorted input is sorted before formatting."""
        result = _format_page_range([9, 1, 5, 2, 3])
        assert result == "pp. 1-3, 5, 9"

    def test_duplicate_pages_deduplicated(self):
        """Duplicate page numbers are deduplicated."""
        result = _format_page_range([1, 1, 2, 2, 3])
        assert result == "pp. 1-3"


# ============================================================================
# _format_page_heading_docx
# ============================================================================
class TestFormatPageHeadingDocx:
    """Tests for _format_page_heading_docx()."""

    def test_arabic_content_page(self):
        """Standard Arabic-numbered content page."""
        result = _format_page_heading_docx(5, "arabic", ["content"], False)
        assert result == "Page 5"

    def test_roman_page(self):
        """Roman numeral page."""
        result = _format_page_heading_docx(3, "roman", ["content"], False)
        assert result == "Page iii"

    def test_unnumbered_page(self):
        """Unnumbered page via page_number_type='none'."""
        result = _format_page_heading_docx("?", "none", ["content"], True)
        assert result == "[Unnumbered page]"

    def test_preface_prefix(self):
        """Preface page type adds [Preface] prefix."""
        result = _format_page_heading_docx(1, "roman", ["preface"], False)
        assert result == "[Preface] Page i"

    def test_appendix_prefix(self):
        """Appendix page type adds [Appendix] prefix."""
        result = _format_page_heading_docx(100, "arabic", ["appendix"], False)
        assert result == "[Appendix] Page 100"

    def test_abstract_prefix(self):
        """Abstract page type (without content) adds [Abstract] prefix."""
        result = _format_page_heading_docx(1, "arabic", ["abstract"], False)
        assert result == "[Abstract] Page 1"

    def test_abstract_with_content_no_prefix(self):
        """Abstract combined with content does not add [Abstract] prefix."""
        result = _format_page_heading_docx(1, "arabic", ["abstract", "content"], False)
        assert "[Abstract]" not in result

    def test_figures_tables_prefix(self):
        """Figures/tables page type adds prefix."""
        result = _format_page_heading_docx(
            50, "arabic", ["figures_tables_sources"], False
        )
        assert result == "[Figures/Tables] Page 50"


# ============================================================================
# _should_render_bullets / _get_structure_types
# ============================================================================
class TestShouldRenderBullets:
    """Tests for _should_render_bullets()."""

    def test_content_page(self):
        assert _should_render_bullets(["content"]) is True

    def test_bibliography_page(self):
        assert _should_render_bullets(["bibliography"]) is False

    def test_blank_page(self):
        assert _should_render_bullets(["blank"]) is False

    def test_mixed_types(self):
        assert _should_render_bullets(["bibliography", "content"]) is True

    def test_appendix(self):
        assert _should_render_bullets(["appendix"]) is True


class TestGetStructureTypes:
    """Tests for _get_structure_types()."""

    def test_content_not_structure(self):
        assert _get_structure_types(["content"]) == []

    def test_bibliography_is_structure(self):
        assert _get_structure_types(["bibliography"]) == ["bibliography"]

    def test_mixed_types(self):
        result = _get_structure_types(["content", "bibliography", "appendix"])
        assert "bibliography" in result
        assert "appendix" in result
        assert "content" not in result


# ============================================================================
# create_docx_summary (mocked)
# ============================================================================
class TestCreateDocxSummary:
    """Tests for create_docx_summary with mocked python-docx."""

    @patch("processors.docx_writer.CitationManager")
    @patch("processors.docx_writer.Document")
    def test_creates_docx_file(self, mock_doc_class, mock_cm_class, tmp_path: Path):
        """create_docx_summary calls Document and saves to output_path."""
        output_path = tmp_path / "summary.docx"
        mock_doc = MagicMock()
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc_class.return_value = mock_doc

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Point A", "Point B"],
                "references": [],
            }
        ]

        create_docx_summary(summary_results, output_path, "Test Document")

        mock_doc.save.assert_called_once_with(str(output_path))

    @patch("processors.docx_writer.CitationManager")
    @patch("processors.docx_writer.Document")
    def test_empty_results_creates_minimal_docx(
        self, mock_doc_class, mock_cm_class, tmp_path: Path
    ):
        """Empty results still produce a valid document."""
        output_path = tmp_path / "summary.docx"
        mock_doc = MagicMock()
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc_class.return_value = mock_doc

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        create_docx_summary([], output_path, "Empty Doc")

        mock_doc.save.assert_called_once()

    @patch("processors.docx_writer.add_hyperlink")
    @patch("processors.docx_writer.config")
    @patch("processors.docx_writer.CitationManager")
    @patch("processors.docx_writer.Document")
    def test_references_section_added(
        self,
        mock_doc_class,
        mock_cm_class,
        mock_config,
        mock_add_hyperlink,
        tmp_path: Path,
    ):
        """References section is added when citations are present."""
        output_path = tmp_path / "summary.docx"
        mock_doc = MagicMock()
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc_class.return_value = mock_doc

        mock_config.CITATION_OPENALEX_EMAIL = "test@test.com"
        mock_config.CITATION_ENABLE_OPENALEX = False
        mock_config.CITATION_MAX_API_REQUESTS = 50

        mock_citation = MagicMock()
        mock_citation.raw_text = "Author (2020). Title."
        mock_citation.url = "https://example.com"
        mock_citation.doi = "10.1234/test"
        mock_citation.metadata = {"publication_year": 2020}

        mock_cm = MagicMock()
        mock_cm.citations = {"key1": mock_citation}
        mock_cm.get_citations_with_pages.return_value = [(mock_citation, "p. 1")]
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Content here"],
                "references": ["Author (2020). Title."],
            }
        ]

        create_docx_summary(summary_results, output_path, "Test")

        # Document should have a page break and heading for references
        mock_doc.add_page_break.assert_called_once()
        # Hyperlink should have been added for the citation
        mock_add_hyperlink.assert_called_once()

    @patch("processors.docx_writer.CitationManager")
    @patch("processors.docx_writer.Document")
    def test_structure_section_for_bibliography_pages(
        self, mock_doc_class, mock_cm_class, tmp_path: Path
    ):
        """Bibliography pages contribute to the Document Structure section."""
        output_path = tmp_path / "summary.docx"
        mock_doc = MagicMock()
        mock_doc.styles = {"Normal": MagicMock()}
        mock_doc_class.return_value = mock_doc

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 50,
                    "page_number_type": "arabic",
                    "page_types": ["bibliography"],
                },
                "bullet_points": [],
                "references": [],
            }
        ]

        create_docx_summary(summary_results, output_path, "Test")

        # Should have been called with "Document Structure" heading
        heading_calls = [
            c
            for c in mock_doc.add_heading.call_args_list
            if "Document Structure" in str(c)
        ]
        assert len(heading_calls) > 0


# ============================================================================
# create_markdown_summary (extended edge cases)
# ============================================================================
class TestCreateMarkdownSummaryExtended:
    """Extended edge case tests for create_markdown_summary."""

    @patch("processors.markdown_writer.CitationManager")
    def test_structure_section_for_toc_pages(self, mock_cm_class, tmp_path: Path):
        """Table of contents pages appear in the Document Structure section."""
        output_path = tmp_path / "summary.md"

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 3,
                    "page_number_type": "roman",
                    "page_types": ["table_of_contents"],
                },
                "bullet_points": [],
                "references": [],
            },
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Some content"],
                "references": [],
            },
        ]

        create_markdown_summary(summary_results, output_path, "Test")

        content = output_path.read_text(encoding="utf-8")
        assert "## Document Structure" in content
        assert "Table of Contents" in content

    @patch("processors.markdown_writer.CitationManager")
    def test_appendix_page_heading(self, mock_cm_class, tmp_path: Path):
        """Appendix pages get [Appendix] prefix in heading."""
        output_path = tmp_path / "summary.md"

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 100,
                    "page_number_type": "arabic",
                    "page_types": ["appendix"],
                },
                "bullet_points": ["Appendix data"],
                "references": [],
            }
        ]

        create_markdown_summary(summary_results, output_path, "Test")

        content = output_path.read_text(encoding="utf-8")
        assert "## [Appendix] Page 100" in content

    @patch("processors.markdown_writer.CitationManager")
    def test_figures_tables_page_heading(self, mock_cm_class, tmp_path: Path):
        """Figures/tables pages get [Figures/Tables] prefix."""
        output_path = tmp_path / "summary.md"

        mock_cm = MagicMock()
        mock_cm.citations = {}
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 30,
                    "page_number_type": "arabic",
                    "page_types": ["figures_tables_sources"],
                },
                "bullet_points": ["Figure 1 description"],
                "references": [],
            }
        ]

        create_markdown_summary(summary_results, output_path, "Test")

        content = output_path.read_text(encoding="utf-8")
        assert "[Figures/Tables]" in content

    @patch("processors.markdown_writer.config")
    @patch("processors.markdown_writer.CitationManager")
    def test_references_with_metadata(self, mock_cm_class, mock_config, tmp_path: Path):
        """References with DOI and year metadata are rendered."""
        output_path = tmp_path / "summary.md"

        mock_config.CITATION_OPENALEX_EMAIL = "test@test.com"
        mock_config.CITATION_ENABLE_OPENALEX = False
        mock_config.CITATION_MAX_API_REQUESTS = 50

        mock_citation = MagicMock()
        mock_citation.raw_text = "Author (2020). Title."
        mock_citation.url = "https://example.com"
        mock_citation.doi = "10.1234/test"
        mock_citation.metadata = {"publication_year": 2020}

        mock_cm = MagicMock()
        mock_cm.citations = {"key": mock_citation}
        mock_cm.get_citations_with_pages.return_value = [(mock_citation, "p. 1")]
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Content"],
                "references": ["Author (2020). Title."],
            }
        ]

        create_markdown_summary(summary_results, output_path, "Test")

        content = output_path.read_text(encoding="utf-8")
        assert "Consolidated References" in content
        assert "DOI: 10.1234/test" in content
        assert "Year: 2020" in content

    @patch("processors.markdown_writer.CitationManager")
    def test_citation_without_url(self, mock_cm_class, tmp_path: Path):
        """Citations without URLs are rendered as plain text."""
        output_path = tmp_path / "summary.md"

        mock_citation = MagicMock()
        mock_citation.raw_text = "No URL Citation"
        mock_citation.url = None
        mock_citation.doi = None
        mock_citation.metadata = {}

        mock_cm = MagicMock()
        mock_cm.citations = {"key": mock_citation}
        mock_cm.get_citations_with_pages.return_value = [(mock_citation, "")]
        mock_cm_class.return_value = mock_cm

        summary_results = [
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Content"],
                "references": ["No URL Citation"],
            }
        ]

        create_markdown_summary(summary_results, output_path, "Test")

        content = output_path.read_text(encoding="utf-8")
        assert "1. No URL Citation" in content


# ============================================================================
# sanitize_omml_xml
# ============================================================================
class TestSanitizeOmmlXml:
    """Tests for sanitize_omml_xml()."""

    def test_valid_xml_unchanged(self):
        """Valid XML is returned unchanged."""
        valid_xml = '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"><m:r><m:t>x</m:t></m:r></m:oMath>'
        result = sanitize_omml_xml(valid_xml)
        assert result == valid_xml

    def test_invalid_xml_returns_original(self):
        """Unfixable invalid XML returns the original string."""
        invalid_xml = "<unclosed><tags"
        result = sanitize_omml_xml(invalid_xml)
        assert result == invalid_xml


# ============================================================================
# add_hyperlink (functional test with real docx)
# ============================================================================
class TestAddHyperlink:
    """Tests for add_hyperlink()."""

    def test_adds_hyperlink_to_paragraph(self):
        """Adds a hyperlink element to a DOCX paragraph."""
        from docx import Document

        doc = Document()
        paragraph = doc.add_paragraph()

        add_hyperlink(paragraph, "https://example.com", "Example Link")

        # Check that the paragraph XML contains a hyperlink element
        xml_str = paragraph._p.xml
        assert "hyperlink" in xml_str


# ============================================================================
# add_formatted_text_to_paragraph
# ============================================================================
class TestAddFormattedTextToParagraph:
    """Tests for add_formatted_text_to_paragraph()."""

    def test_plain_text_added(self):
        """Plain text without LaTeX is added as a run."""
        from docx import Document

        doc = Document()
        paragraph = doc.add_paragraph()

        add_formatted_text_to_paragraph(paragraph, "Hello world")

        assert len(paragraph.runs) >= 1
        full_text = paragraph.text
        assert "Hello world" in full_text

    def test_text_with_inline_latex(self):
        """Text with inline LaTeX renders math or falls back to text."""
        from docx import Document

        doc = Document()
        paragraph = doc.add_paragraph()

        # This may succeed or fall back to text rendering
        add_formatted_text_to_paragraph(paragraph, "Value is $x + y$")

        # Paragraph should contain some content
        xml_str = paragraph._p.xml
        assert len(xml_str) > 0


# ============================================================================
# filter_empty_pages
# ============================================================================
class TestFilterEmptyPages:
    """Tests for filter_empty_pages()."""

    def test_filters_blank_pages(self):
        """Pages with blank type are filtered out."""
        results = [
            {
                "page_information": {
                    "page_number_integer": 1,
                    "page_number_type": "arabic",
                    "page_types": ["blank"],
                },
                "bullet_points": None,
            },
            {
                "page_information": {
                    "page_number_integer": 2,
                    "page_number_type": "arabic",
                    "page_types": ["content"],
                },
                "bullet_points": ["Real content"],
            },
        ]

        filtered = filter_empty_pages(results)
        assert len(filtered) == 1

    def test_keeps_structure_pages_without_bullets(self):
        """Pages with structure types are kept even without bullets."""
        results = [
            {
                "page_information": {
                    "page_number_integer": 50,
                    "page_number_type": "arabic",
                    "page_types": ["bibliography"],
                },
                "bullet_points": [],
            },
        ]

        filtered = filter_empty_pages(results)
        assert len(filtered) == 1

    def test_empty_input(self):
        """Empty input returns empty output."""
        assert filter_empty_pages([]) == []

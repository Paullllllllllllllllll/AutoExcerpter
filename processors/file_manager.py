"""File management utilities for transcription and summary outputs.

This module handles all file I/O operations for the AutoExcerpter application:

1. **DOCX Summary Creation**: 
   - Generates formatted Word documents with summaries
   - Manages citation deduplication and consolidation
   - Adds hyperlinks to citations with metadata
   - Applies consistent formatting (headings, bullets, spacing)

2. **Transcription Text Output**:
   - Writes raw transcriptions to text files
   - Includes metadata (timestamps, statistics, source info)

3. **JSON Logging**:
   - Initializes structured log files for API results
   - Appends entries during processing
   - Finalizes logs with proper JSON formatting

4. **Content Filtering**:
   - Filters out empty/error pages from summaries
   - Validates semantic content

5. **XML Safety**:
   - Sanitizes text for safe DOCX/XML output
   - Removes invalid control characters

All functions handle errors gracefully and provide detailed logging.
"""

from __future__ import annotations

import json
import re
import threading
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Optional
from xml.etree import ElementTree as ET

import mathml2omml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from latex2mathml.converter import convert as latex_to_mathml

from modules import app_config as config
from modules.concurrency_helper import (
    get_api_concurrency,
    get_api_timeout,
    get_service_tier,
)
from modules.constants import (
    MATH_NAMESPACE,
    TITLE_HEADING_LEVEL,
    PAGE_HEADING_LEVEL,
    TITLE_SPACE_AFTER_PT,
    PAGE_HEADING_SPACE_BEFORE_PT,
    PAGE_HEADING_SPACE_AFTER_PT,
    BULLET_SPACE_AFTER_PT,
    REF_INDENT_PT,
    BULLET_INDENT_PT,
    ERROR_MARKERS,
)
from modules.logger import setup_logger
from modules.roman_numerals import int_to_roman
from processors.citation_manager import CitationManager

logger = setup_logger(__name__)


_LOG_HANDLES: dict[Path, tuple[Any, threading.Lock]] = {}
_LOG_HANDLES_GUARD = threading.Lock()


def _get_log_handle(log_path: Path):
    key = log_path
    with _LOG_HANDLES_GUARD:
        existing = _LOG_HANDLES.get(key)
        if existing is not None:
            return existing
        handle = key.open("a", encoding="utf-8")
        lock = threading.Lock()
        _LOG_HANDLES[key] = (handle, lock)
        return handle, lock


def _close_log_handle(log_path: Path) -> None:
    key = log_path
    with _LOG_HANDLES_GUARD:
        existing = _LOG_HANDLES.pop(key, None)
    if existing is None:
        return
    handle, _lock = existing
    try:
        handle.close()
    except Exception:
        pass


# ============================================================================
# Text Sanitization Functions
# ============================================================================
def sanitize_for_xml(text: Optional[str]) -> str:
    """Return XML-safe text for DOCX output by removing control characters."""
    if not text:
        return ""

    # Remove control characters (except tab \x09, newline \x0A, and carriage return \x0D)
    sanitized = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)
    return sanitized


def parse_latex_in_text(text: str) -> list[tuple[str, str]]:
    """Parse text and split it into segments of regular text and LaTeX formulas.
    
    Handles:
    - Display math: $$...$$ (processed first to avoid conflicts)
    - Inline math: $...$ (single dollar, non-greedy)
    - Escaped dollar signs: \\$ are preserved as literal $
    - Multi-line formulas (DOTALL mode)
    - Nested braces within formulas
    
    Returns:
        List of (content, type) tuples where type is 'text', 'latex_display', or 'latex_inline'.
    """
    if not text:
        return [('', 'text')]
    
    # Placeholder for escaped dollar signs to protect them during parsing
    ESCAPED_DOLLAR_PLACEHOLDER = '\x00ESCAPED_DOLLAR\x00'
    
    # Protect escaped dollar signs
    protected_text = text.replace('\\$', ESCAPED_DOLLAR_PLACEHOLDER)
    
    segments: list[tuple[str, str]] = []
    
    # Pattern for display math $$...$$ (process first - greedy on delimiters, non-greedy on content)
    # Uses negative lookbehind to avoid matching escaped $
    display_pattern = r'\$\$(.+?)\$\$'
    
    # Pattern for inline math $...$ (single $, non-greedy, must not be empty)
    # Negative lookbehind/lookahead to avoid matching $$ or empty $$
    inline_pattern = r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)'
    
    # First pass: extract display math $$...$$
    last_end = 0
    temp_segments: list[tuple[str, str, int, int]] = []  # (content, type, start, end)
    
    for match in re.finditer(display_pattern, protected_text, re.DOTALL):
        temp_segments.append((match.group(1), 'latex_display', match.start(), match.end()))
    
    # Second pass: extract inline math $...$ from non-display regions
    # Build list of display math ranges to exclude
    display_ranges = [(s[2], s[3]) for s in temp_segments]
    
    def in_display_range(pos: int) -> bool:
        return any(start <= pos < end for start, end in display_ranges)
    
    for match in re.finditer(inline_pattern, protected_text, re.DOTALL):
        # Only add if not inside a display math region
        if not in_display_range(match.start()) and not in_display_range(match.end()):
            temp_segments.append((match.group(1), 'latex_inline', match.start(), match.end()))
    
    # Sort all segments by position
    temp_segments.sort(key=lambda x: x[2])
    
    # Build final segments list with text between formulas
    last_end = 0
    for content, seg_type, start, end in temp_segments:
        # Add text before this formula
        if start > last_end:
            text_segment = protected_text[last_end:start]
            # Restore escaped dollars in text
            text_segment = text_segment.replace(ESCAPED_DOLLAR_PLACEHOLDER, '$')
            if text_segment:
                segments.append((text_segment, 'text'))
        
        # Add the formula (restore any escaped dollars that might be in it)
        restored_content = content.replace(ESCAPED_DOLLAR_PLACEHOLDER, '$')
        segments.append((restored_content, seg_type))
        last_end = end
    
    # Add remaining text after last formula
    if last_end < len(protected_text):
        remaining = protected_text[last_end:]
        remaining = remaining.replace(ESCAPED_DOLLAR_PLACEHOLDER, '$')
        if remaining:
            segments.append((remaining, 'text'))
    
    # If no segments found, return original text
    if not segments:
        return [(text, 'text')]
    
    return segments


def normalize_latex_whitespace(latex_code: str) -> str:
    """Normalize whitespace in LaTeX formulas for consistent processing."""
    # Strip leading/trailing whitespace
    normalized = latex_code.strip()
    # Collapse multiple spaces to single space
    normalized = re.sub(r' +', ' ', normalized)
    # Remove spaces around operators that don't need them
    normalized = re.sub(r'\s*([=+\-*/^_{}])\s*', r'\1', normalized)
    return normalized


def simplify_problematic_latex(latex_code: str) -> tuple[str, list[str]]:
    """Simplify LaTeX code by removing constructs known to cause mathml2omml issues.
    
    This function applies progressive simplifications to LaTeX code that is known
    to cause issues with the latex2mathml -> mathml2omml conversion pipeline.
    
    Returns:
        Tuple of (simplified_latex, list_of_applied_simplifications)
    """
    simplified = latex_code
    applied = []
    
    # 1. Normalize whitespace first
    original_len = len(simplified)
    simplified = normalize_latex_whitespace(simplified)
    if len(simplified) != original_len:
        applied.append('whitespace normalization')
    
    # 2. Replace logical operators (not supported by latex2mathml)
    logical_ops = [
        ('\\land', '\\wedge'),
        ('\\lor', '\\vee'),
        ('\\lnot', '\\neg'),
        ('\\iff', '\\Leftrightarrow'),
        ('\\implies', '\\Rightarrow'),
    ]
    for old, new in logical_ops:
        if old in simplified:
            simplified = simplified.replace(old, new)
            applied.append(f'logical operator ({old} -> {new})')
    
    # 3. Replace text commands that cause issues
    text_commands = [
        (r'\\text\{([^}]*)\}', r'\\mathrm{\1}', 'text -> mathrm'),
        (r'\\textrm\{([^}]*)\}', r'\\mathrm{\1}', 'textrm -> mathrm'),
        (r'\\textit\{([^}]*)\}', r'\\mathit{\1}', 'textit -> mathit'),
        (r'\\textbf\{([^}]*)\}', r'\\mathbf{\1}', 'textbf -> mathbf'),
    ]
    for pattern, replacement, desc in text_commands:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(desc)
    
    # 4. Replace accent macros that cause groupChr issues
    accent_replacements = [
        (r'\\bar\{([^}]+)\}', r'\\overline{\1}', 'bar -> overline'),
        (r'\\tilde\{([^}]+)\}', r'\\widetilde{\1}', 'tilde -> widetilde'),
        (r'\\hat\{([^}]+)\}', r'\\widehat{\1}', 'hat -> widehat'),
        (r'\\vec\{([^}]+)\}', r'\\overrightarrow{\1}', 'vec -> overrightarrow'),
        (r'\\dot\{([^}]+)\}', r'\1', 'dot removed'),
        (r'\\ddot\{([^}]+)\}', r'\1', 'ddot removed'),
    ]
    for pattern, replacement, desc in accent_replacements:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f'accent ({desc})')
    
    # 5. Replace delimiter size commands
    delimiter_sizes = [
        '\\Big', '\\big', '\\bigg', '\\Bigg', 
        '\\Bigl', '\\Bigr', '\\bigl', '\\bigr', 
        '\\biggl', '\\biggr', '\\Biggl', '\\Biggr',
        '\\bigm', '\\Bigm'
    ]
    delimiter_chars = ['(', ')', '[', ']', '|', '\\{', '\\}', '\\|', '.']
    
    delimiter_simplified = False
    for size in delimiter_sizes:
        for delim in delimiter_chars:
            target = f'{size}{delim}'
            if target in simplified:
                # Replace sized delimiter with plain delimiter
                plain_delim = delim if delim not in ['\\{', '\\}', '\\|'] else delim[1:]
                simplified = simplified.replace(target, plain_delim if delim != '.' else '')
                delimiter_simplified = True
    
    if delimiter_simplified:
        applied.append('delimiter sizing removed')
    
    # 6. Replace auto-sizing delimiters
    if '\\left' in simplified or '\\right' in simplified:
        # Handle \left. and \right. (invisible delimiters)
        simplified = simplified.replace('\\left.', '')
        simplified = simplified.replace('\\right.', '')
        # Handle normal delimiters
        simplified = re.sub(r'\\left\s*([(\[|])', r'\1', simplified)
        simplified = re.sub(r'\\right\s*([)\]|])', r'\1', simplified)
        simplified = re.sub(r'\\left\s*\\([{|}])', r'\\\1', simplified)
        simplified = re.sub(r'\\right\s*\\([{|}])', r'\\\1', simplified)
        simplified = simplified.replace('\\middle|', '|')
        simplified = simplified.replace('\\middle', '')
        applied.append('auto-sizing delimiters (left/right) removed')
    
    # 7. Replace overbrace/underbrace (complex structures)
    brace_patterns = [
        (r'\\overbrace\{([^}]+)\}\^?\{?[^}]*\}?', r'\1', 'overbrace'),
        (r'\\underbrace\{([^}]+)\}_?\{?[^}]*\}?', r'\1', 'underbrace'),
        (r'\\overleftarrow\{([^}]+)\}', r'\1', 'overleftarrow'),
        (r'\\overrightarrow\{([^}]+)\}', r'\\vec{\1}', 'overrightarrow'),
        (r'\\underleftarrow\{([^}]+)\}', r'\1', 'underleftarrow'),
        (r'\\underrightarrow\{([^}]+)\}', r'\1', 'underrightarrow'),
    ]
    for pattern, replacement, desc in brace_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f'{desc} simplified')
    
    # 8. Replace \phantom and \hphantom (invisible spacing)
    phantom_patterns = [
        (r'\\phantom\{[^}]*\}', '', 'phantom'),
        (r'\\hphantom\{[^}]*\}', '', 'hphantom'),
        (r'\\vphantom\{[^}]*\}', '', 'vphantom'),
    ]
    for pattern, replacement, desc in phantom_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f'{desc} removed')
    
    # 9. Replace problematic spacing commands (literal strings)
    spacing_literals = [
        ('\\,', ' ', 'thin space'),
        ('\\;', ' ', 'thick space'),
        ('\\!', '', 'negative thin space'),
        ('\\quad', ' ', 'quad'),
        ('\\qquad', '  ', 'qquad'),
    ]
    for old, new, desc in spacing_literals:
        if old in simplified:
            simplified = simplified.replace(old, new)
            applied.append(f'{desc} normalized')
    
    # Spacing commands with arguments (regex patterns)
    spacing_patterns = [
        (r'\\hspace\{[^}]*\}', ' ', 'hspace'),
        (r'\\vspace\{[^}]*\}', '', 'vspace'),
        (r'\\mspace\{[^}]*\}', ' ', 'mspace'),
    ]
    for pattern, replacement, desc in spacing_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f'{desc} normalized')
    
    # 10. Handle \operatorname (convert to mathrm for compatibility)
    if '\\operatorname' in simplified:
        simplified = re.sub(r'\\operatorname\{([^}]+)\}', r'\\mathrm{\1}', simplified)
        applied.append('operatorname -> mathrm')
    
    # 11. Handle common unsupported environments
    env_patterns = [
        (r'\\begin\{align\*?\}(.+?)\\end\{align\*?\}', r'\1', 'align env'),
        (r'\\begin\{gather\*?\}(.+?)\\end\{gather\*?\}', r'\1', 'gather env'),
        (r'\\begin\{equation\*?\}(.+?)\\end\{equation\*?\}', r'\1', 'equation env'),
        (r'\\begin\{split\}(.+?)\\end\{split\}', r'\1', 'split env'),
        (r'\\begin\{cases\}(.+?)\\end\{cases\}', r'\1', 'cases env'),
    ]
    for pattern, replacement, desc in env_patterns:
        if re.search(pattern, simplified, re.DOTALL):
            simplified = re.sub(pattern, replacement, simplified, flags=re.DOTALL)
            applied.append(f'{desc} unwrapped')
    
    # 12. Clean up alignment markers from environments
    if '&' in simplified or '\\\\' in simplified:
        # Remove alignment markers
        simplified = simplified.replace('&', ' ')
        # Replace line breaks with spaces
        simplified = simplified.replace('\\\\', ' ')
        applied.append('alignment markers cleaned')
    
    # 13. Handle \limits and \nolimits
    if '\\limits' in simplified or '\\nolimits' in simplified:
        simplified = simplified.replace('\\limits', '')
        simplified = simplified.replace('\\nolimits', '')
        applied.append('limits modifiers removed')
    
    # 14. Final whitespace cleanup
    simplified = re.sub(r'\s+', ' ', simplified).strip()
    
    return simplified, applied


def sanitize_omml_xml(omml_markup: str) -> str:
    """Attempt to fix common XML issues in OMML markup generated by mathml2omml."""
    try:
        # Try to parse as-is first
        ET.fromstring(omml_markup)
        return omml_markup
    except ET.ParseError as e:
        error_msg = str(e)
        
        # Attempt to fix groupChr tag mismatch issues
        # This is a known bug in mathml2omml where groupChrPr tags are not properly closed
        if 'groupChr' in error_msg:
            # Strategy 1: Find groupChrPr tags without closing tags before the parent groupChr closes
            # Match: <m:groupChrPr...> followed by content that doesn't contain </m:groupChrPr>
            # but does contain </m:groupChr>
            fixed_markup = omml_markup
            
            # Find all groupChrPr opening tags and check if they're properly closed
            pattern = r'<m:groupChrPr[^>]*>'
            positions = [(m.start(), m.end()) for m in re.finditer(pattern, fixed_markup)]
            
            # Process from end to start to maintain positions
            for start, end in reversed(positions):
                # Extract the section after this groupChrPr tag
                after_tag = fixed_markup[end:]
                
                # Check if there's a closing </m:groupChrPr> before the next </m:groupChr>
                close_prop = after_tag.find('</m:groupChrPr>')
                close_group = after_tag.find('</m:groupChr>')
                
                # If groupChr closes before groupChrPr, we need to insert the closing tag
                if close_group != -1 and (close_prop == -1 or close_prop > close_group):
                    # Insert </m:groupChrPr> right before </m:groupChr>
                    insert_pos = end + close_group
                    fixed_markup = fixed_markup[:insert_pos] + '</m:groupChrPr>' + fixed_markup[insert_pos:]
            
            # Verify the fix worked
            try:
                ET.fromstring(fixed_markup)
                logger.info("Successfully sanitized OMML with groupChr fix")
                return fixed_markup
            except ET.ParseError as parse_err:
                logger.debug("First sanitization attempt failed: %s", parse_err)
                
                # Strategy 2: More aggressive - close all unclosed groupChrPr tags
                # before any groupChr closing tag, handling multiple occurrences
                fixed_markup2 = omml_markup
                
                # Iteratively fix each groupChr block
                max_iterations = 10  # Prevent infinite loops
                for _ in range(max_iterations):
                    # Find groupChr blocks with unclosed groupChrPr
                    match = re.search(
                        r'<m:groupChr>\s*<m:groupChrPr[^>]*>(?:(?!</m:groupChrPr>).)*?</m:groupChr>',
                        fixed_markup2,
                        flags=re.DOTALL
                    )
                    
                    if not match:
                        break  # No more unclosed tags
                    
                    # Insert closing tag before the groupChr end
                    block = match.group(0)
                    fixed_block = block.replace('</m:groupChr>', '</m:groupChrPr></m:groupChr>', 1)
                    fixed_markup2 = fixed_markup2[:match.start()] + fixed_block + fixed_markup2[match.end():]
                
                try:
                    ET.fromstring(fixed_markup2)
                    logger.info("Successfully sanitized OMML with iterative groupChr fix")
                    return fixed_markup2
                except ET.ParseError:
                    logger.debug("Iterative sanitization failed, trying fallback")
                    
                    # Strategy 3: Nuclear option - replace all groupChr blocks with simplified structure
                    # This removes the problematic property tags entirely
                    fixed_markup3 = re.sub(
                        r'<m:groupChr>\s*<m:groupChrPr[^>]*>.*?</m:groupChr>',
                        lambda m: m.group(0).replace('<m:groupChrPr', '<!--m:groupChrPr').replace('</m:groupChrPr>', '-->'),
                        omml_markup,
                        flags=re.DOTALL
                    )
                    
                    try:
                        ET.fromstring(fixed_markup3)
                        logger.info("Successfully sanitized OMML by commenting out groupChrPr")
                        return fixed_markup3
                    except ET.ParseError:
                        pass
        
        # If we can't fix it, return original
        return omml_markup


def add_math_to_paragraph(paragraph, latex_code: str) -> None:
    """Add a mathematical formula to a paragraph using OMML for native Word equation rendering."""
    try:
        # Convert LaTeX -> MathML -> OMML
        mathml = latex_to_mathml(latex_code)
        omml_markup = mathml2omml.convert(mathml)

        # Ensure the OMML output has the proper namespace declaration
        if "xmlns:m" not in omml_markup.split("\n", 1)[0]:
            omml_markup = omml_markup.replace(
                "<m:oMath",
                f'<m:oMath xmlns:m="{MATH_NAMESPACE}"',
                1,
            )

        # Wrap in oMathPara for compatibility with Word paragraphs
        if not omml_markup.strip().startswith("<m:oMath"):
            omml_markup = f'<m:oMath xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMath>'

        omml_para = (
            f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMathPara>'
        )
        
        # Try to parse the XML
        try:
            omml_element = parse_xml(omml_para)
        except Exception as parse_error:
            # Attempt to sanitize the OMML XML
            logger.warning("XML parsing failed, attempting to sanitize OMML: %s", parse_error)
            sanitized_omml = sanitize_omml_xml(omml_markup)
            sanitized_para = f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{sanitized_omml}</m:oMathPara>'
            omml_element = parse_xml(sanitized_para)

        # Append native equation to the paragraph
        paragraph._p.append(omml_element)
        
    except Exception as exc:
        # If all conversion attempts fail, try a simplified version
        try:
            simplified_latex, simplifications = simplify_problematic_latex(latex_code)
            
            if simplified_latex != latex_code:
                logger.info("Attempting conversion with simplified LaTeX: %s", ', '.join(simplifications))
                mathml = latex_to_mathml(simplified_latex)
                omml_markup = mathml2omml.convert(mathml)
                
                if "xmlns:m" not in omml_markup.split("\n", 1)[0]:
                    omml_markup = omml_markup.replace(
                        "<m:oMath",
                        f'<m:oMath xmlns:m="{MATH_NAMESPACE}"',
                        1,
                    )
                
                if not omml_markup.strip().startswith("<m:oMath"):
                    omml_markup = f'<m:oMath xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMath>'
                
                omml_para = f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMathPara>'
                omml_element = parse_xml(omml_para)
                paragraph._p.append(omml_element)
                logger.info("Successfully converted simplified LaTeX")
                return
        except:
            pass
        
        # Final fallback: display as formatted text
        logger.warning("Failed to convert LaTeX to MathML: %s. Displaying as text.", exc)
        logger.warning("Problematic LaTeX code: %s", latex_code[:200])  # Log first 200 chars
        run = paragraph.add_run(f" {latex_code} ")
        run.font.name = 'Cambria Math'
        run.italic = True


def add_formatted_text_to_paragraph(paragraph, text: str) -> None:
    """Add text to a paragraph, parsing and rendering LaTeX formulas.
    
    Handles both display math ($$...$$) and inline math ($...$) with appropriate
    rendering for each type.
    """
    segments = parse_latex_in_text(text)
    
    for content, segment_type in segments:
        if segment_type in ('latex', 'latex_display', 'latex_inline'):
            # Both display and inline math are rendered the same way in Word
            # The distinction could be used for different styling in future
            add_math_to_paragraph(paragraph, content)
        else:
            paragraph.add_run(sanitize_for_xml(content))


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a hyperlink to a paragraph in a DOCX document."""
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run object (a wrapper over a text element)
    new_run = OxmlElement('w:r')
    
    # Set the run's text
    rPr = OxmlElement('w:rPr')
    
    # Add run properties for hyperlink styling
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    
    # Add the actual text
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)


def _extract_summary_payload(result: dict[str, Any]) -> dict[str, Any]:
    """Return the summary payload dict, handling both flat and nested formats.
    
    Flat structure (preferred): page_information, bullet_points, references at top level.
    Legacy nested: summary.summary containing those fields.
    """
    # Check for flat structure first (page_information at top level)
    if "page_information" in result and isinstance(result.get("page_information"), dict):
        return result
    
    # Fall back to legacy nested structure
    summary = result.get("summary", {})
    if not isinstance(summary, dict):
        return {}

    nested_summary = summary.get("summary")
    if isinstance(nested_summary, dict):
        return nested_summary
    return summary


def _page_information(summary_data: dict[str, Any]) -> dict[str, Any]:
    """Extract normalized page information from summary payload.
    
    Returns dict with:
        - page_number_integer: The numeric page number (or "?" if null/missing)
        - page_number_type: 'roman', 'arabic', or 'none'
        - page_types: List of content classifications (content, bibliography, etc.)
        - is_unnumbered: Boolean flag derived from page_number_type == 'none' or null integer
    """
    page_info = summary_data.get("page_information", {})
    
    # Handle dict format (preferred) - must be non-empty dict
    if isinstance(page_info, dict) and page_info:
        page_int = page_info.get("page_number_integer")
        page_num_type = page_info.get("page_number_type", "arabic")
        
        # Handle both page_types (array) and legacy page_type (string)
        page_types = page_info.get("page_types")
        if page_types is None:
            # Fallback to legacy page_type field
            legacy_type = page_info.get("page_type", "content")
            page_types = [legacy_type] if legacy_type else ["content"]
        elif isinstance(page_types, str):
            page_types = [page_types]
        elif not isinstance(page_types, list) or not page_types:
            page_types = ["content"]
        
        # Derive unnumbered status from page_number_type or null page_number_integer
        is_unnumbered = (page_num_type == "none" or page_int is None)
        if is_unnumbered:
            page_num_type = "none"
            page_int = "?"
        
        return {
            "page_number_integer": page_int,
            "page_number_type": page_num_type,
            "page_types": page_types,
            "is_unnumbered": is_unnumbered,
        }
    
    # Fallback to page field
    page_val = summary_data.get("page", "?")
    return {
        "page_number_integer": page_val,
        "page_number_type": "arabic",
        "page_types": ["content"],
        "is_unnumbered": False,
    }


# Page types that should have bullet points extracted (summarizable prose)
PAGE_TYPES_WITH_BULLETS = {"content", "abstract", "preface", "appendix", "figures_tables_sources"}

# Page types shown in Document Structure section (ordered by typical document position)
STRUCTURE_PAGE_TYPE_ORDER = [
    "title_page",           # Front matter
    "copyright",
    "abstract",             # Often first content page
    "table_of_contents",
    "preface",
    "figures_tables_sources",
    "appendix",             # Back matter
    "bibliography",
    "index",
    "other",
    # "blank" excluded - never shown in structure
]

# Human-readable labels for page types
PAGE_TYPE_LABELS = {
    "content": "Content",
    "preface": "Preface",
    "appendix": "Appendix",
    "figures_tables_sources": "Figures, Tables & Sources",
    "table_of_contents": "Table of Contents",
    "bibliography": "Bibliography",
    "title_page": "Title Page",
    "index": "Index",
    "blank": "Blank Pages",
    "abstract": "Abstract",
    "copyright": "Copyright",
    "other": "Other",
}


def _should_render_bullets(page_types: list[str]) -> bool:
    """Check if page should have bullet points rendered based on its types."""
    return bool(set(page_types) & PAGE_TYPES_WITH_BULLETS)


def _get_structure_types(page_types: list[str]) -> list[str]:
    """Return page types that should appear in Document Structure section."""
    return [pt for pt in page_types if pt in STRUCTURE_PAGE_TYPE_ORDER]


def _is_meaningful_summary(summary_data: dict[str, Any]) -> bool:
    """Check if a summary is meaningful based on page_types and content.
    
    Pages with bullet-point types (content, abstract, preface, appendix, figures_tables_sources)
    are meaningful if they have non-error bullet points.
    
    Pages with structure types (bibliography, TOC, etc.) are meaningful for 
    page range display in Document Structure section.
    """
    page_info = _page_information(summary_data)
    page_types = page_info.get("page_types", ["content"])
    
    # Blank pages are never meaningful (only if blank is the sole type)
    if page_types == ["blank"]:
        return False
    
    # Check if page has bullet-point types
    has_bullet_types = _should_render_bullets(page_types)
    
    if has_bullet_types:
        bullet_points = summary_data.get("bullet_points") or []
        if bullet_points:
            # Check for error markers
            if len(bullet_points) == 1:
                marker_text = bullet_points[0].strip().lower()
                if any(marker in marker_text for marker in ERROR_MARKERS):
                    # Has bullet types but error content - still meaningful for structure
                    return bool(_get_structure_types(page_types))
            return True
        # No bullets but has structure types - still meaningful for structure section
        return bool(_get_structure_types(page_types))
    
    # No bullet types - meaningful if has structure types for page range display
    return bool(_get_structure_types(page_types))


def _format_page_range(pages: list[int]) -> str:
    """Format a list of page numbers as a compact range string.
    
    Examples:
        [1, 2, 3, 5, 7, 8, 9] -> "pp. 1-3, 5, 7-9"
        [5] -> "p. 5"
    """
    if not pages:
        return ""
    
    pages = sorted(set(pages))
    if len(pages) == 1:
        return f"p. {pages[0]}"
    
    ranges = []
    start = pages[0]
    end = pages[0]
    
    for page in pages[1:]:
        if page == end + 1:
            end = page
        else:
            if start == end:
                ranges.append(str(start))
            else:
                ranges.append(f"{start}-{end}")
            start = end = page
    
    # Add last range
    if start == end:
        ranges.append(str(start))
    else:
        ranges.append(f"{start}-{end}")
    
    return f"pp. {', '.join(ranges)}"


def _format_page_heading_docx(
    page_number: Any, page_number_type: str, page_types: list[str], is_unnumbered: bool
) -> str:
    """Format page heading for DOCX output based on page_types and numbering."""
    # Add page type prefix for non-content pages with bullets
    type_prefix = ""
    if "abstract" in page_types and "content" not in page_types:
        type_prefix = "[Abstract] "
    elif "preface" in page_types:
        type_prefix = "[Preface] "
    elif "appendix" in page_types:
        type_prefix = "[Appendix] "
    elif "figures_tables_sources" in page_types:
        type_prefix = "[Figures/Tables] "
    
    if page_number_type == "roman" and isinstance(page_number, int):
        roman_str = int_to_roman(page_number)
        return f"{type_prefix}Page {roman_str}"
    elif page_number_type == "none" or is_unnumbered:
        return f"{type_prefix}[Unnumbered page]"
    else:
        return f"{type_prefix}Page {page_number}"


def create_docx_summary(
    summary_results: list[dict[str, Any]], output_path: Path, document_name: str
) -> None:
    """Create a DOCX summary document from structured summary results.
    
    Output order:
    1. Title + Metadata
    2. Document Structure (page type overview)
    3. Content Summaries (in document order)
    4. Consolidated References
    """
    filtered_results = filter_empty_pages(summary_results)
    if len(filtered_results) < len(summary_results):
        logger.info(
            "Filtered out %s pages with no useful content",
            len(summary_results) - len(filtered_results),
        )

    # Initialize citation manager for the document
    citation_manager = CitationManager(polite_pool_email=config.CITATION_OPENALEX_EMAIL)
    
    # Collect page numbers by page_type for Document Structure section
    # Use ordered dict based on STRUCTURE_PAGE_TYPE_ORDER
    page_type_pages: dict[str, list[int]] = {pt: [] for pt in STRUCTURE_PAGE_TYPE_ORDER}
    
    # First pass: collect structure info and citations from all pages
    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_types = page_info["page_types"]
        references = summary_payload.get("references") or []
        
        # Collect citations from ALL page types
        if references and isinstance(page_number, int):
            citation_manager.add_citations(references, page_number)
        
        # Collect page numbers for structure section (all applicable types)
        if isinstance(page_number, int):
            for pt in _get_structure_types(page_types):
                page_type_pages[pt].append(page_number)
    
    document = Document()

    # Configure default Normal style for compact spacing
    normal_style = document.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)

    # === SECTION 1: Title ===
    title = document.add_heading(f"Summary of {sanitize_for_xml(document_name)}", TITLE_HEADING_LEVEL)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)

    # Metadata line
    content_pages = sum(
        1 for r in filtered_results 
        if _should_render_bullets(_page_information(_extract_summary_payload(r)).get("page_types", ["content"]))
    )
    metadata = "Processed: %s | Content pages: %s | Total pages: %s" % (
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        content_pages,
        len(filtered_results),
    )
    meta_paragraph = document.add_paragraph(metadata)
    meta_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_paragraph.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)

    # === SECTION 2: Document Structure (at the beginning) ===
    has_structure_info = any(pages for pages in page_type_pages.values())
    if has_structure_info:
        struct_heading = document.add_heading("Document Structure", PAGE_HEADING_LEVEL)
        struct_heading.paragraph_format.space_before = Pt(PAGE_HEADING_SPACE_BEFORE_PT)
        struct_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)
        
        # Add page ranges in STRUCTURE_PAGE_TYPE_ORDER
        for pt in STRUCTURE_PAGE_TYPE_ORDER:
            pages = page_type_pages.get(pt, [])
            if pages:
                label = PAGE_TYPE_LABELS.get(pt, pt.replace("_", " ").title())
                page_range = _format_page_range(pages)
                struct_para = document.add_paragraph()
                struct_para.paragraph_format.space_before = Pt(0)
                struct_para.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
                label_run = struct_para.add_run(f"{label}: ")
                label_run.bold = True
                struct_para.add_run(page_range)

    # === SECTION 3: Content Summaries (in document order) ===
    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_number_type = page_info["page_number_type"]
        page_types = page_info["page_types"]
        bullet_points = summary_payload.get("bullet_points") or []

        # Render summary for pages with bullet-point types AND actual bullets
        if _should_render_bullets(page_types) and bullet_points:
            heading_text = _format_page_heading_docx(
                page_number, page_number_type, page_types, page_info["is_unnumbered"]
            )
            page_heading = document.add_heading(heading_text, PAGE_HEADING_LEVEL)
            page_heading.paragraph_format.space_before = Pt(PAGE_HEADING_SPACE_BEFORE_PT)
            page_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)

            # Bullet points - use Word's native bullet list style
            for point in bullet_points:
                paragraph = document.add_paragraph(style='List Bullet')
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
                add_formatted_text_to_paragraph(paragraph, point)

    # Add consolidated references section at the end of the document
    if citation_manager.citations:
        logger.info("Processing %d unique citations for consolidated references section", 
                   len(citation_manager.citations))
        
        # Enrich citations with metadata from OpenAlex API (if enabled)
        if config.CITATION_ENABLE_OPENALEX:
            citation_manager.enrich_with_metadata(max_requests=config.CITATION_MAX_API_REQUESTS)
        else:
            logger.info("OpenAlex enrichment disabled - skipping metadata lookup")
        
        # Add references heading
        document.add_page_break()
        ref_main_heading = document.add_heading("Consolidated References", TITLE_HEADING_LEVEL)
        ref_main_heading.paragraph_format.space_before = Pt(PAGE_HEADING_SPACE_BEFORE_PT)
        ref_main_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)
        
        # Add note about the references
        note_text = ("The following references were extracted from the document and consolidated. "
                    "Duplicate citations have been merged, showing all pages where each citation appears. "
                    "Where available, hyperlinks provide access to extended metadata via OpenAlex.")
        note_paragraph = document.add_paragraph(note_text)
        note_paragraph.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)
        
        # Get sorted citations with page information
        citations_with_pages = citation_manager.get_citations_with_pages()
        
        for idx, (citation, page_range_str) in enumerate(citations_with_pages, start=1):
            # Create citation paragraph
            ref_paragraph = document.add_paragraph()
            ref_paragraph.paragraph_format.space_before = Pt(0)
            ref_paragraph.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
            ref_paragraph.paragraph_format.left_indent = Pt(REF_INDENT_PT)
            ref_paragraph.paragraph_format.first_line_indent = Pt(-REF_INDENT_PT)
            
            # Add citation number
            num_run = ref_paragraph.add_run(f"[{idx}] ")
            num_run.bold = True
            
            # Add citation text
            citation_text = sanitize_for_xml(citation.raw_text)
            
            # If we have metadata with a URL, add as hyperlink
            if citation.url:
                # Add citation text as hyperlink
                add_hyperlink(ref_paragraph, citation.url, citation_text)
            else:
                ref_paragraph.add_run(citation_text)
            
            # Add page information
            if page_range_str:
                page_run = ref_paragraph.add_run(f" ({page_range_str})")
                page_run.italic = True
            
            # Add metadata note if available
            if citation.metadata:
                meta_info_parts = []
                if citation.doi:
                    meta_info_parts.append(f"DOI: {citation.doi}")
                if citation.metadata.get('publication_year'):
                    meta_info_parts.append(f"Year: {citation.metadata['publication_year']}")
                
                if meta_info_parts:
                    meta_run = ref_paragraph.add_run(f" [{', '.join(meta_info_parts)}]")
                    meta_run.font.size = Pt(9)
                    meta_run.italic = True

    # === SECTION 4: Save document ===
    document.save(str(output_path))
    logger.info("Summary document saved to %s", output_path)


def _format_page_heading_md(
    page_number: Any, page_number_type: str, page_types: list[str], is_unnumbered: bool
) -> str:
    """Format page heading for markdown output based on page_types and numbering."""
    # Add page type prefix for non-content pages with bullets
    type_prefix = ""
    if "abstract" in page_types and "content" not in page_types:
        type_prefix = "[Abstract] "
    elif "preface" in page_types:
        type_prefix = "[Preface] "
    elif "appendix" in page_types:
        type_prefix = "[Appendix] "
    elif "figures_tables_sources" in page_types:
        type_prefix = "[Figures/Tables] "
    
    if page_number_type == "roman" and isinstance(page_number, int):
        roman_str = int_to_roman(page_number)
        return f"## {type_prefix}Page {roman_str}"
    elif page_number_type == "none" or is_unnumbered:
        return f"## {type_prefix}[Unnumbered page]"
    else:
        return f"## {type_prefix}Page {page_number}"


def create_markdown_summary(
    summary_results: list[dict[str, Any]], output_path: Path, document_name: str
) -> None:
    """Create a Markdown summary document from structured summary results.
    
    Output order:
    1. Title + Metadata
    2. Document Structure (page type overview)
    3. Content Summaries (in document order)
    4. Consolidated References
    
    Args:
        summary_results: List of summary result dictionaries from the API.
        output_path: Path where the markdown file will be written.
        document_name: Name of the source document for the title.
    """
    filtered_results = filter_empty_pages(summary_results)
    if len(filtered_results) < len(summary_results):
        logger.info(
            "Filtered out %s pages with no useful content",
            len(summary_results) - len(filtered_results),
        )
    
    # Initialize citation manager for the document
    citation_manager = CitationManager(polite_pool_email=config.CITATION_OPENALEX_EMAIL)
    
    # Collect page numbers by page_type for Document Structure section
    page_type_pages: dict[str, list[int]] = {pt: [] for pt in STRUCTURE_PAGE_TYPE_ORDER}
    
    # First pass: collect structure info and citations from all pages
    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_types = page_info["page_types"]
        references = summary_payload.get("references") or []
        
        # Collect citations from ALL page types
        if references and isinstance(page_number, int):
            citation_manager.add_citations(references, page_number)
        
        # Collect page numbers for structure section (all applicable types)
        if isinstance(page_number, int):
            for pt in _get_structure_types(page_types):
                page_type_pages[pt].append(page_number)
    
    lines: list[str] = []
    
    # === SECTION 1: Title ===
    lines.append(f"# Summary of {sanitize_for_xml(document_name)}")
    lines.append("")
    
    # Metadata
    content_pages = sum(
        1 for r in filtered_results 
        if _should_render_bullets(_page_information(_extract_summary_payload(r)).get("page_types", ["content"]))
    )
    lines.append(
        f"*Processed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
        f"Content pages: {content_pages} | Total pages: {len(filtered_results)}*"
    )
    lines.append("")
    
    # === SECTION 2: Document Structure (at the beginning) ===
    has_structure_info = any(pages for pages in page_type_pages.values())
    if has_structure_info:
        lines.append("## Document Structure")
        lines.append("")
        
        # Add page ranges in STRUCTURE_PAGE_TYPE_ORDER
        for pt in STRUCTURE_PAGE_TYPE_ORDER:
            pages = page_type_pages.get(pt, [])
            if pages:
                label = PAGE_TYPE_LABELS.get(pt, pt.replace("_", " ").title())
                page_range = _format_page_range(pages)
                lines.append(f"- **{label}**: {page_range}")
        
        lines.append("")
        lines.append("---")
        lines.append("")
    
    # === SECTION 3: Content Summaries (in document order) ===
    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_number_type = page_info["page_number_type"]
        page_types = page_info["page_types"]
        bullet_points = summary_payload.get("bullet_points") or []
        
        # Render summary for pages with bullet-point types AND actual bullets
        if _should_render_bullets(page_types) and bullet_points:
            lines.append(_format_page_heading_md(
                page_number, page_number_type, page_types, page_info["is_unnumbered"]
            ))
            lines.append("")
            
            # Bullet points (LaTeX formulas preserved as-is for markdown compatibility)
            for point in bullet_points:
                sanitized_point = sanitize_for_xml(point)
                lines.append(f"- {sanitized_point}")
            
            lines.append("")
    
    # Add consolidated references section at the end
    if citation_manager.citations:
        logger.info(
            "Processing %d unique citations for consolidated references section",
            len(citation_manager.citations),
        )
        
        # Enrich citations with metadata from OpenAlex API (if enabled)
        if config.CITATION_ENABLE_OPENALEX:
            citation_manager.enrich_with_metadata(max_requests=config.CITATION_MAX_API_REQUESTS)
        else:
            logger.info("OpenAlex enrichment disabled - skipping metadata lookup")
        
        lines.append("---")
        lines.append("")
        lines.append("## Consolidated References")
        lines.append("")
        lines.append(
            "*The following references were extracted from the document and consolidated. "
            "Duplicate citations have been merged, showing all pages where each citation appears. "
            "Where available, hyperlinks provide access to extended metadata via OpenAlex.*"
        )
        lines.append("")
        
        # Get sorted citations with page information
        citations_with_pages = citation_manager.get_citations_with_pages()
        
        for idx, (citation, page_range_str) in enumerate(citations_with_pages, start=1):
            citation_text = sanitize_for_xml(citation.raw_text)
            
            # Build citation line with optional hyperlink
            if citation.url:
                line = f"{idx}. [{citation_text}]({citation.url})"
            else:
                line = f"{idx}. {citation_text}"
            
            # Add page range
            if page_range_str:
                line += f" *({page_range_str})*"
            
            # Add metadata if available
            if citation.metadata:
                meta_parts = []
                if citation.doi:
                    meta_parts.append(f"DOI: {citation.doi}")
                if citation.metadata.get("publication_year"):
                    meta_parts.append(f"Year: {citation.metadata['publication_year']}")
                if meta_parts:
                    line += f" *[{', '.join(meta_parts)}]*"
            
            lines.append(line)
    
    # === SECTION 4: Write file ===
    output_path.write_text("\n".join(lines), encoding="utf-8")
    logger.info("Markdown summary saved to %s", output_path)


def write_transcription_to_text(
    transcription_results: list[dict[str, Any]],
    output_path: Path,
    document_name: str,
    item_type: str,
    total_elapsed_time: float,
    source_path: Path,
) -> bool:
    """Persist transcription output as a text file alongside basic metadata."""
    elapsed_str = str(timedelta(seconds=int(total_elapsed_time)))
    successes = sum(1 for result in transcription_results if "error" not in result)
    failures = len(transcription_results) - successes

    try:
        with output_path.open("w", encoding="utf-8") as file_handle:
            file_handle.write(f"# Transcription of: {document_name}\n")
            file_handle.write(f"# Source Path: {source_path}\n")
            file_handle.write(f"# Type: {item_type}\n")
            file_handle.write(
                f"# Processed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            )
            file_handle.write(f"# Total images processed: {len(transcription_results)}\n")
            file_handle.write(f"# Successfully transcribed: {successes}\n")
            file_handle.write(f"# Failed items: {failures}\n")
            file_handle.write(
                f"# Total processing time for this item: {elapsed_str}\n\n"
            )

            for index, result in enumerate(transcription_results):
                file_handle.write(
                    result.get("transcription", "[ERROR] Transcription data missing")
                )
                if index < len(transcription_results) - 1:
                    file_handle.write("\n")

        logger.info("Transcription text file saved: %s", output_path)
        return True
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Error writing transcription to text file %s: %s", output_path, exc)
        return False


def initialize_log_file(
    log_path: Path,
    item_name: str,
    input_path: str,
    input_type: str,
    total_images: int,
    model_name: str,
    extraction_dpi: Optional[int] = None,
    concurrency_limit: Optional[int] = None,
) -> bool:
    """Create the per-item log file header as the start of a JSON array."""
    # Determine if this is an OpenAI model for flex processing metadata
    is_openai_model = model_name.startswith(("gpt-", "o1", "o3", "o4"))
    default_concurrency, _ = get_api_concurrency()
    service_tier = get_service_tier() if is_openai_model else "N/A"
    configuration = {
        "concurrent_requests": concurrency_limit if concurrency_limit is not None else default_concurrency,
        "api_timeout_seconds": get_api_timeout(),
        "model_name": model_name,
        "extraction_dpi": extraction_dpi if extraction_dpi is not None else "N/A",
        "service_tier": service_tier,
    }

    payload = {
        "input_item_name": item_name,
        "input_item_path": input_path,
        "input_type": input_type,
        "processing_start_time": datetime.now().isoformat(),
        "total_images": total_images,
        "configuration": configuration,
    }

    try:
        _close_log_handle(log_path)
        with log_path.open("w", encoding="utf-8") as log_file:
            log_file.write("[\n")  # Start JSON array
            json.dump(payload, log_file)
        return True
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("Failed to initialize log file %s: %s", log_path, exc)
        return False


def append_to_log(log_path: Path, entry: dict[str, Any]) -> bool:
    """Append a JSON entry to the log file array (comma-separated)."""
    try:
        log_file, lock = _get_log_handle(log_path)
        with lock:
            log_file.write(",\n")  # Add comma separator
            json.dump(entry, log_file)
        return True
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("Failed to write to log file %s: %s", log_path, exc)
        return False


def finalize_log_file(log_path: Path) -> bool:
    """Finalize the log file by closing the JSON array."""
    try:
        log_file, lock = _get_log_handle(log_path)
        with lock:
            log_file.write("\n]")  # Close JSON array
        _close_log_handle(log_path)
        return True
    except Exception as exc:  # pylint: disable=broad-except
        logger.warning("Failed to finalize log file %s: %s", log_path, exc)
        return False


def filter_empty_pages(summary_results: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Remove entries without meaningful summary content."""
    filtered: list[dict[str, Any]] = []
    for result in summary_results:
        summary_payload = _extract_summary_payload(result)
        if summary_payload and _is_meaningful_summary(summary_payload):
            filtered.append(result)
    return filtered


# ============================================================================
# Public API
# ============================================================================
__all__ = [
    "create_docx_summary",
    "create_markdown_summary",
    "write_transcription_to_text",
    "initialize_log_file",
    "append_to_log",
    "finalize_log_file",
    "filter_empty_pages",
    "sanitize_for_xml",
]
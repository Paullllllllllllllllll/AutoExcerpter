"""DOCX summary document creation for AutoExcerpter.

Generates formatted Word documents with structured summaries, including:
- LaTeX formula parsing and native Word equation rendering (OMML)
- Citation deduplication with OpenAlex hyperlinks
- Document structure overview with page ranges
- XML sanitization for safe DOCX output
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import mathml2omml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from latex2mathml.converter import convert as latex_to_mathml

from modules import app_config as config
from modules.constants import (
    MATH_NAMESPACE,
    TITLE_HEADING_LEVEL,
    PAGE_HEADING_LEVEL,
    TITLE_SPACE_AFTER_PT,
    PAGE_HEADING_SPACE_BEFORE_PT,
    PAGE_HEADING_SPACE_AFTER_PT,
    BULLET_SPACE_AFTER_PT,
    REF_INDENT_PT,
)
from modules.logger import setup_logger
from modules.roman_numerals import int_to_roman
from processors.citation_manager import CitationManager
from processors.file_manager import (
    sanitize_for_xml,
    _extract_summary_payload,
    _page_information,
    _should_render_bullets,
    _get_structure_types,
    _format_page_range,
    filter_empty_pages,
    STRUCTURE_PAGE_TYPE_ORDER,
    PAGE_TYPE_LABELS,
)

logger = setup_logger(__name__)


def parse_latex_in_text(text: str) -> list[tuple[str, str]]:
    """Parse text and split it into segments of regular text and LaTeX formulas.

    Handles:
    - Display math: $$...$$ (processed first to avoid conflicts)
    - Inline math: $...$ (single dollar, non-greedy)
    - Escaped dollar signs: \\$ are preserved as literal $
    - Multi-line formulas (DOTALL mode)
    - Nested braces within formulas

    Returns:
        List of (content, type) tuples where type is 'text', 'latex_display', or 'latex_inline'.
    """
    if not text:
        return [("", "text")]

    # Placeholder for escaped dollar signs to protect them during parsing
    ESCAPED_DOLLAR_PLACEHOLDER = "\x00ESCAPED_DOLLAR\x00"

    # Protect escaped dollar signs
    protected_text = text.replace("\\$", ESCAPED_DOLLAR_PLACEHOLDER)

    segments: list[tuple[str, str]] = []

    # Pattern for display math $$...$$ (process first - greedy on delimiters, non-greedy on content)
    display_pattern = r"\$\$(.+?)\$\$"

    # Pattern for inline math $...$ (single $, non-greedy, must not be empty)
    inline_pattern = r"(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)"

    # First pass: extract display math $$...$$
    temp_segments: list[tuple[str, str, int, int]] = []  # (content, type, start, end)

    for match in re.finditer(display_pattern, protected_text, re.DOTALL):
        temp_segments.append(
            (match.group(1), "latex_display", match.start(), match.end())
        )

    # Second pass: extract inline math $...$ from non-display regions
    display_ranges = [(s[2], s[3]) for s in temp_segments]

    def in_display_range(pos: int) -> bool:
        return any(start <= pos < end for start, end in display_ranges)

    for match in re.finditer(inline_pattern, protected_text, re.DOTALL):
        if not in_display_range(match.start()) and not in_display_range(match.end()):
            temp_segments.append(
                (match.group(1), "latex_inline", match.start(), match.end())
            )

    # Sort all segments by position
    temp_segments.sort(key=lambda x: x[2])

    # Build final segments list with text between formulas
    last_end = 0
    for content, seg_type, start, end in temp_segments:
        if start > last_end:
            text_segment = protected_text[last_end:start]
            text_segment = text_segment.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
            if text_segment:
                segments.append((text_segment, "text"))

        restored_content = content.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
        segments.append((restored_content, seg_type))
        last_end = end

    # Add remaining text after last formula
    if last_end < len(protected_text):
        remaining = protected_text[last_end:]
        remaining = remaining.replace(ESCAPED_DOLLAR_PLACEHOLDER, "$")
        if remaining:
            segments.append((remaining, "text"))

    if not segments:
        return [(text, "text")]

    return segments


def normalize_latex_whitespace(latex_code: str) -> str:
    """Normalize whitespace in LaTeX formulas for consistent processing."""
    normalized = latex_code.strip()
    normalized = re.sub(r" +", " ", normalized)
    normalized = re.sub(r"\s*([=+\-*/^_{}])\s*", r"\1", normalized)
    return normalized


def simplify_problematic_latex(latex_code: str) -> tuple[str, list[str]]:
    """Simplify LaTeX code by removing constructs known to cause mathml2omml issues.

    Returns:
        Tuple of (simplified_latex, list_of_applied_simplifications)
    """
    simplified = latex_code
    applied = []

    # 1. Normalize whitespace first
    original_len = len(simplified)
    simplified = normalize_latex_whitespace(simplified)
    if len(simplified) != original_len:
        applied.append("whitespace normalization")

    # 2. Replace logical operators (not supported by latex2mathml)
    logical_ops = [
        ("\\land", "\\wedge"),
        ("\\lor", "\\vee"),
        ("\\lnot", "\\neg"),
        ("\\iff", "\\Leftrightarrow"),
        ("\\implies", "\\Rightarrow"),
    ]
    for old, new in logical_ops:
        if old in simplified:
            simplified = simplified.replace(old, new)
            applied.append(f"logical operator ({old} -> {new})")

    # 3. Replace text commands that cause issues
    text_commands = [
        (r"\\text\{([^}]*)\}", r"\\mathrm{\1}", "text -> mathrm"),
        (r"\\textrm\{([^}]*)\}", r"\\mathrm{\1}", "textrm -> mathrm"),
        (r"\\textit\{([^}]*)\}", r"\\mathit{\1}", "textit -> mathit"),
        (r"\\textbf\{([^}]*)\}", r"\\mathbf{\1}", "textbf -> mathbf"),
    ]
    for pattern, replacement, desc in text_commands:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(desc)

    # 4. Replace accent macros that cause groupChr issues
    accent_replacements = [
        (r"\\bar\{([^}]+)\}", r"\\overline{\1}", "bar -> overline"),
        (r"\\tilde\{([^}]+)\}", r"\\widetilde{\1}", "tilde -> widetilde"),
        (r"\\hat\{([^}]+)\}", r"\\widehat{\1}", "hat -> widehat"),
        (r"\\vec\{([^}]+)\}", r"\\overrightarrow{\1}", "vec -> overrightarrow"),
        (r"\\dot\{([^}]+)\}", r"\1", "dot removed"),
        (r"\\ddot\{([^}]+)\}", r"\1", "ddot removed"),
    ]
    for pattern, replacement, desc in accent_replacements:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f"accent ({desc})")

    # 5. Replace delimiter size commands
    delimiter_sizes = [
        "\\Big",
        "\\big",
        "\\bigg",
        "\\Bigg",
        "\\Bigl",
        "\\Bigr",
        "\\bigl",
        "\\bigr",
        "\\biggl",
        "\\biggr",
        "\\Biggl",
        "\\Biggr",
        "\\bigm",
        "\\Bigm",
    ]
    delimiter_chars = ["(", ")", "[", "]", "|", "\\{", "\\}", "\\|", "."]

    delimiter_simplified = False
    for size in delimiter_sizes:
        for delim in delimiter_chars:
            target = f"{size}{delim}"
            if target in simplified:
                plain_delim = delim if delim not in ["\\{", "\\}", "\\|"] else delim[1:]
                simplified = simplified.replace(
                    target, plain_delim if delim != "." else ""
                )
                delimiter_simplified = True

    if delimiter_simplified:
        applied.append("delimiter sizing removed")

    # 6. Replace auto-sizing delimiters
    if "\\left" in simplified or "\\right" in simplified:
        simplified = simplified.replace("\\left.", "")
        simplified = simplified.replace("\\right.", "")
        simplified = re.sub(r"\\left\s*([(\[|])", r"\1", simplified)
        simplified = re.sub(r"\\right\s*([)\]|])", r"\1", simplified)
        simplified = re.sub(r"\\left\s*\\([{|}])", r"\\\1", simplified)
        simplified = re.sub(r"\\right\s*\\([{|}])", r"\\\1", simplified)
        simplified = simplified.replace("\\middle|", "|")
        simplified = simplified.replace("\\middle", "")
        applied.append("auto-sizing delimiters (left/right) removed")

    # 7. Replace overbrace/underbrace (complex structures)
    brace_patterns = [
        (r"\\overbrace\{([^}]+)\}\^?\{?[^}]*\}?", r"\1", "overbrace"),
        (r"\\underbrace\{([^}]+)\}_?\{?[^}]*\}?", r"\1", "underbrace"),
        (r"\\overleftarrow\{([^}]+)\}", r"\1", "overleftarrow"),
        (r"\\overrightarrow\{([^}]+)\}", r"\\vec{\1}", "overrightarrow"),
        (r"\\underleftarrow\{([^}]+)\}", r"\1", "underleftarrow"),
        (r"\\underrightarrow\{([^}]+)\}", r"\1", "underrightarrow"),
    ]
    for pattern, replacement, desc in brace_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f"{desc} simplified")

    # 8. Replace \phantom and \hphantom (invisible spacing)
    phantom_patterns = [
        (r"\\phantom\{[^}]*\}", "", "phantom"),
        (r"\\hphantom\{[^}]*\}", "", "hphantom"),
        (r"\\vphantom\{[^}]*\}", "", "vphantom"),
    ]
    for pattern, replacement, desc in phantom_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f"{desc} removed")

    # 9. Replace problematic spacing commands (literal strings)
    spacing_literals = [
        ("\\,", " ", "thin space"),
        ("\\;", " ", "thick space"),
        ("\\!", "", "negative thin space"),
        ("\\quad", " ", "quad"),
        ("\\qquad", "  ", "qquad"),
    ]
    for old, new, desc in spacing_literals:
        if old in simplified:
            simplified = simplified.replace(old, new)
            applied.append(f"{desc} normalized")

    # Spacing commands with arguments (regex patterns)
    spacing_patterns = [
        (r"\\hspace\{[^}]*\}", " ", "hspace"),
        (r"\\vspace\{[^}]*\}", "", "vspace"),
        (r"\\mspace\{[^}]*\}", " ", "mspace"),
    ]
    for pattern, replacement, desc in spacing_patterns:
        if re.search(pattern, simplified):
            simplified = re.sub(pattern, replacement, simplified)
            applied.append(f"{desc} normalized")

    # 10. Handle \operatorname (convert to mathrm for compatibility)
    if "\\operatorname" in simplified:
        simplified = re.sub(r"\\operatorname\{([^}]+)\}", r"\\mathrm{\1}", simplified)
        applied.append("operatorname -> mathrm")

    # 11. Handle common unsupported environments
    env_patterns = [
        (r"\\begin\{align\*?\}(.+?)\\end\{align\*?\}", r"\1", "align env"),
        (r"\\begin\{gather\*?\}(.+?)\\end\{gather\*?\}", r"\1", "gather env"),
        (r"\\begin\{equation\*?\}(.+?)\\end\{equation\*?\}", r"\1", "equation env"),
        (r"\\begin\{split\}(.+?)\\end\{split\}", r"\1", "split env"),
        (r"\\begin\{cases\}(.+?)\\end\{cases\}", r"\1", "cases env"),
    ]
    for pattern, replacement, desc in env_patterns:
        if re.search(pattern, simplified, re.DOTALL):
            simplified = re.sub(pattern, replacement, simplified, flags=re.DOTALL)
            applied.append(f"{desc} unwrapped")

    # 12. Clean up alignment markers from environments
    if "&" in simplified or "\\\\" in simplified:
        simplified = simplified.replace("&", " ")
        simplified = simplified.replace("\\\\", " ")
        applied.append("alignment markers cleaned")

    # 13. Handle \limits and \nolimits
    if "\\limits" in simplified or "\\nolimits" in simplified:
        simplified = simplified.replace("\\limits", "")
        simplified = simplified.replace("\\nolimits", "")
        applied.append("limits modifiers removed")

    # 14. Final whitespace cleanup
    simplified = re.sub(r"\s+", " ", simplified).strip()

    return simplified, applied


def sanitize_omml_xml(omml_markup: str) -> str:
    """Attempt to fix common XML issues in OMML markup generated by mathml2omml."""
    try:
        ET.fromstring(omml_markup)
        return omml_markup
    except ET.ParseError as e:
        error_msg = str(e)

        if "groupChr" in error_msg:
            fixed_markup = omml_markup

            pattern = r"<m:groupChrPr[^>]*>"
            positions = [
                (m.start(), m.end()) for m in re.finditer(pattern, fixed_markup)
            ]

            for start, end in reversed(positions):
                after_tag = fixed_markup[end:]
                close_prop = after_tag.find("</m:groupChrPr>")
                close_group = after_tag.find("</m:groupChr>")

                if close_group != -1 and (close_prop == -1 or close_prop > close_group):
                    insert_pos = end + close_group
                    fixed_markup = (
                        fixed_markup[:insert_pos]
                        + "</m:groupChrPr>"
                        + fixed_markup[insert_pos:]
                    )

            try:
                ET.fromstring(fixed_markup)
                logger.info("Successfully sanitized OMML with groupChr fix")
                return fixed_markup
            except ET.ParseError as parse_err:
                logger.debug("First sanitization attempt failed: %s", parse_err)

                fixed_markup2 = omml_markup
                max_iterations = 10
                for _ in range(max_iterations):
                    match = re.search(
                        r"<m:groupChr>\s*<m:groupChrPr[^>]*>(?:(?!</m:groupChrPr>).)*?</m:groupChr>",
                        fixed_markup2,
                        flags=re.DOTALL,
                    )
                    if not match:
                        break
                    block = match.group(0)
                    fixed_block = block.replace(
                        "</m:groupChr>", "</m:groupChrPr></m:groupChr>", 1
                    )
                    fixed_markup2 = (
                        fixed_markup2[: match.start()]
                        + fixed_block
                        + fixed_markup2[match.end() :]
                    )

                try:
                    ET.fromstring(fixed_markup2)
                    logger.info(
                        "Successfully sanitized OMML with iterative groupChr fix"
                    )
                    return fixed_markup2
                except ET.ParseError:
                    logger.debug("Iterative sanitization failed, trying fallback")

                    fixed_markup3 = re.sub(
                        r"<m:groupChr>\s*<m:groupChrPr[^>]*>.*?</m:groupChr>",
                        lambda m: m.group(0)
                        .replace("<m:groupChrPr", "<!--m:groupChrPr")
                        .replace("</m:groupChrPr>", "-->"),
                        omml_markup,
                        flags=re.DOTALL,
                    )

                    try:
                        ET.fromstring(fixed_markup3)
                        logger.info(
                            "Successfully sanitized OMML by commenting out groupChrPr"
                        )
                        return fixed_markup3
                    except ET.ParseError:
                        pass

        return omml_markup


def add_math_to_paragraph(paragraph, latex_code: str) -> None:
    """Add a mathematical formula to a paragraph using OMML for native Word equation rendering."""
    try:
        mathml = latex_to_mathml(latex_code)
        omml_markup = mathml2omml.convert(mathml)

        if "xmlns:m" not in omml_markup.split("\n", 1)[0]:
            omml_markup = omml_markup.replace(
                "<m:oMath",
                f'<m:oMath xmlns:m="{MATH_NAMESPACE}"',
                1,
            )

        if not omml_markup.strip().startswith("<m:oMath"):
            omml_markup = f'<m:oMath xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMath>'

        omml_para = (
            f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMathPara>'
        )

        try:
            omml_element = parse_xml(omml_para)
        except (ET.ParseError, ValueError) as parse_error:
            logger.warning(
                "XML parsing failed, attempting to sanitize OMML: %s", parse_error
            )
            sanitized_omml = sanitize_omml_xml(omml_markup)
            sanitized_para = f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{sanitized_omml}</m:oMathPara>'
            omml_element = parse_xml(sanitized_para)

        paragraph._p.append(omml_element)

    except Exception as exc:
        try:
            simplified_latex, simplifications = simplify_problematic_latex(latex_code)

            if simplified_latex != latex_code:
                logger.info(
                    "Attempting conversion with simplified LaTeX: %s",
                    ", ".join(simplifications),
                )
                mathml = latex_to_mathml(simplified_latex)
                omml_markup = mathml2omml.convert(mathml)

                if "xmlns:m" not in omml_markup.split("\n", 1)[0]:
                    omml_markup = omml_markup.replace(
                        "<m:oMath",
                        f'<m:oMath xmlns:m="{MATH_NAMESPACE}"',
                        1,
                    )

                if not omml_markup.strip().startswith("<m:oMath"):
                    omml_markup = (
                        f'<m:oMath xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMath>'
                    )

                omml_para = f'<m:oMathPara xmlns:m="{MATH_NAMESPACE}">{omml_markup}</m:oMathPara>'
                omml_element = parse_xml(omml_para)
                paragraph._p.append(omml_element)
                logger.info("Successfully converted simplified LaTeX")
                return
        except Exception:
            pass

        logger.warning(
            "Failed to convert LaTeX to MathML: %s. Displaying as text.", exc
        )
        logger.warning("Problematic LaTeX code: %s", latex_code[:200])
        run = paragraph.add_run(f" {latex_code} ")
        run.font.name = "Cambria Math"
        run.italic = True


def add_formatted_text_to_paragraph(paragraph, text: str) -> None:
    """Add text to a paragraph, parsing and rendering LaTeX formulas."""
    segments = parse_latex_in_text(text)

    for content, segment_type in segments:
        if segment_type in ("latex", "latex_display", "latex_inline"):
            add_math_to_paragraph(paragraph, content)
        else:
            paragraph.add_run(sanitize_for_xml(content))


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a hyperlink to a paragraph in a DOCX document."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def _format_page_heading_docx(
    page_number: Any, page_number_type: str, page_types: list[str], is_unnumbered: bool
) -> str:
    """Format page heading for DOCX output based on page_types and numbering."""
    type_prefix = ""
    if "abstract" in page_types and "content" not in page_types:
        type_prefix = "[Abstract] "
    elif "preface" in page_types:
        type_prefix = "[Preface] "
    elif "appendix" in page_types:
        type_prefix = "[Appendix] "
    elif "figures_tables_sources" in page_types:
        type_prefix = "[Figures/Tables] "

    if page_number_type == "roman" and isinstance(page_number, int):
        roman_str = int_to_roman(page_number)
        return f"{type_prefix}Page {roman_str}"
    elif page_number_type == "none" or is_unnumbered:
        return f"{type_prefix}[Unnumbered page]"
    else:
        return f"{type_prefix}Page {page_number}"


def create_docx_summary(
    summary_results: list[dict[str, Any]], output_path: Path, document_name: str
) -> None:
    """Create a DOCX summary document from structured summary results.

    Output order:
    1. Title + Metadata
    2. Document Structure (page type overview)
    3. Content Summaries (in document order)
    4. Consolidated References
    """
    filtered_results = filter_empty_pages(summary_results)
    if len(filtered_results) < len(summary_results):
        logger.info(
            "Filtered out %s pages with no useful content",
            len(summary_results) - len(filtered_results),
        )

    citation_manager = CitationManager(polite_pool_email=config.CITATION_OPENALEX_EMAIL)

    page_type_pages: dict[str, list[int]] = {pt: [] for pt in STRUCTURE_PAGE_TYPE_ORDER}

    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_types = page_info["page_types"]
        references = summary_payload.get("references") or []

        if references and isinstance(page_number, int):
            citation_manager.add_citations(references, page_number)

        if isinstance(page_number, int):
            for pt in _get_structure_types(page_types):
                page_type_pages[pt].append(page_number)

    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)

    # === SECTION 1: Title ===
    title = document.add_heading(
        f"Summary of {sanitize_for_xml(document_name)}", TITLE_HEADING_LEVEL
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)

    content_pages = sum(
        1
        for r in filtered_results
        if _should_render_bullets(
            _page_information(_extract_summary_payload(r)).get(
                "page_types", ["content"]
            )
        )
    )
    metadata = "Processed: %s | Content pages: %s | Total pages: %s" % (
        datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        content_pages,
        len(filtered_results),
    )
    meta_paragraph = document.add_paragraph(metadata)
    meta_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    meta_paragraph.paragraph_format.space_after = Pt(TITLE_SPACE_AFTER_PT)

    # === SECTION 2: Document Structure ===
    has_structure_info = any(pages for pages in page_type_pages.values())
    if has_structure_info:
        struct_heading = document.add_heading("Document Structure", PAGE_HEADING_LEVEL)
        struct_heading.paragraph_format.space_before = Pt(PAGE_HEADING_SPACE_BEFORE_PT)
        struct_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)

        for pt in STRUCTURE_PAGE_TYPE_ORDER:
            pages = page_type_pages.get(pt, [])
            if pages:
                label = PAGE_TYPE_LABELS.get(pt, pt.replace("_", " ").title())
                page_range = _format_page_range(pages)
                struct_para = document.add_paragraph()
                struct_para.paragraph_format.space_before = Pt(0)
                struct_para.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
                label_run = struct_para.add_run(f"{label}: ")
                label_run.bold = True
                struct_para.add_run(page_range)

    # === SECTION 3: Content Summaries ===
    for result in filtered_results:
        summary_payload = _extract_summary_payload(result)
        page_info = _page_information(summary_payload)
        page_number = page_info["page_number_integer"]
        page_number_type = page_info["page_number_type"]
        page_types = page_info["page_types"]
        bullet_points = summary_payload.get("bullet_points") or []

        if _should_render_bullets(page_types) and bullet_points:
            heading_text = _format_page_heading_docx(
                page_number, page_number_type, page_types, page_info["is_unnumbered"]
            )
            page_heading = document.add_heading(heading_text, PAGE_HEADING_LEVEL)
            page_heading.paragraph_format.space_before = Pt(
                PAGE_HEADING_SPACE_BEFORE_PT
            )
            page_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)

            for point in bullet_points:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
                add_formatted_text_to_paragraph(paragraph, point)

    # === SECTION 4: Consolidated References ===
    if citation_manager.citations:
        logger.info(
            "Processing %d unique citations for consolidated references section",
            len(citation_manager.citations),
        )

        if config.CITATION_ENABLE_OPENALEX:
            citation_manager.enrich_with_metadata(
                max_requests=config.CITATION_MAX_API_REQUESTS
            )
        else:
            logger.info("OpenAlex enrichment disabled - skipping metadata lookup")

        document.add_page_break()
        ref_main_heading = document.add_heading(
            "Consolidated References", TITLE_HEADING_LEVEL
        )
        ref_main_heading.paragraph_format.space_before = Pt(
            PAGE_HEADING_SPACE_BEFORE_PT
        )
        ref_main_heading.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)

        note_text = (
            "The following references were extracted from the document and consolidated. "
            "Duplicate citations have been merged, showing all pages where each citation appears. "
            "Where available, hyperlinks provide access to extended metadata via OpenAlex."
        )
        note_paragraph = document.add_paragraph(note_text)
        note_paragraph.paragraph_format.space_after = Pt(PAGE_HEADING_SPACE_AFTER_PT)

        citations_with_pages = citation_manager.get_citations_with_pages()

        for idx, (citation, page_range_str) in enumerate(citations_with_pages, start=1):
            ref_paragraph = document.add_paragraph()
            ref_paragraph.paragraph_format.space_before = Pt(0)
            ref_paragraph.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER_PT)
            ref_paragraph.paragraph_format.left_indent = Pt(REF_INDENT_PT)
            ref_paragraph.paragraph_format.first_line_indent = Pt(-REF_INDENT_PT)

            num_run = ref_paragraph.add_run(f"[{idx}] ")
            num_run.bold = True

            citation_text = sanitize_for_xml(citation.raw_text)

            if citation.url:
                add_hyperlink(ref_paragraph, citation.url, citation_text)
            else:
                ref_paragraph.add_run(citation_text)

            if page_range_str:
                page_run = ref_paragraph.add_run(f" ({page_range_str})")
                page_run.italic = True

            if citation.metadata:
                meta_info_parts = []
                if citation.doi:
                    meta_info_parts.append(f"DOI: {citation.doi}")
                if citation.metadata.get("publication_year"):
                    meta_info_parts.append(
                        f"Year: {citation.metadata['publication_year']}"
                    )

                if meta_info_parts:
                    meta_run = ref_paragraph.add_run(f" [{', '.join(meta_info_parts)}]")
                    meta_run.font.size = Pt(9)
                    meta_run.italic = True

    document.save(str(output_path))
    logger.info("Summary document saved to %s", output_path)
